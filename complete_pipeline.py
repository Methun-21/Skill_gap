import streamlit as st
import pandas as pd
import PyPDF2
import docx
import re
import os
import logging
import tempfile
from typing import Dict, List
from io import BytesIO
from datetime import datetime
import json

# --- Page Config ---
st.set_page_config(
    page_title="AI Skill Gap Analyzer",
    page_icon="🧠",
    layout="wide"
)
# --- CUSTOM CSS STYLING ---
st.markdown("""
    <style>
    /* --- Main App Styling --- */
    :root {
        --primary-color: #1e3a8a; /* Deep Blue */
        --secondary-color: #4338ca; /* Indigo */
        --accent-color: #0d9488; /* Teal */
        --text-color: #e2e8f0; /* Light Gray for text */
        --bg-color: #0f172a; /* Slate background */
        --card-bg-color: #1e293b; /* Darker Slate for cards */
    }

    /* Set a dark theme */
    .stApp {
        background-color: var(--bg-color);
        color: var(--text-color);
    }

    /* --- Title and Headers --- */
    h1 {
        color: #f8fafc; /* White */
        background: -webkit-linear-gradient(45deg, #38bdf8, #a78bfa);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    h2, h3 {
        color: #93c5fd; /* Light Blue */
    }

    /* --- Buttons --- */
    .stButton > button {
        border-radius: 20px;
        background-image: linear-gradient(to right, var(--accent-color), #14b8a6);
        color: white;
        border: none;
        padding: 10px 24px;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px 0 rgba(13, 148, 136, 0.4);
    }
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px 0 rgba(13, 148, 136, 0.6);
    }
    .stButton > button:active {
        transform: translateY(0);
    }
    
    /* Style download buttons differently */
    .stDownloadButton > button {
        background-image: linear-gradient(to right, #4338ca, #6d28d9);
        box-shadow: 0 4px 15px 0 rgba(67, 56, 202, 0.4);
    }
    .stDownloadButton > button:hover {
        box-shadow: 0 6px 20px 0 rgba(67, 56, 202, 0.6);
    }


    /* --- Tabs --- */
    button[data-baseweb="tab"][aria-selected="true"] {
        background-image: linear-gradient(to right, var(--secondary-color), #6d28d9);
        color: white;
        border-radius: 8px 8px 0 0;
    }
    
    /* --- Containers and Cards --- */
    [data-testid="stExpander"] {
        background-color: var(--card-bg-color);
        border-radius: 10px;
        border: 1px solid #334155;
    }
    .stDataFrame {
        border-radius: 10px;
    }
    [data-testid="stMetric"] {
        background-color: var(--card-bg-color);
        border-radius: 10px;
        padding: 20px;
        border: 1px solid #334155;
    }

    </style>
""", unsafe_allow_html=True)


# --------------------------- Helper Classes ---------------------------

class DocumentUploader:
    """Handles file upload functionality with validation"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt']
        self.max_file_size = 10 * 1024 * 1024  # 10MB in bytes
    
    def create_upload_interface(self):
        """Create the main upload interface"""
        st.header("Upload Documents 📄")
        
        st.info("Upload your resumes and job descriptions here. The system will process the text and get it ready for analysis. Supported formats: PDF, DOCX, TXT.")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("📄 Resumes")
            resume_files = st.file_uploader(
                "Choose Resume files",
                type=self.supported_formats,
                accept_multiple_files=True,
                help="Upload one or more resume files",
                key="resume_uploader"
            )
            
        with col2:
            st.subheader("💼 Job Descriptions")
            job_files = st.file_uploader(
                "Choose Job Description files",
                type=self.supported_formats,
                accept_multiple_files=True,
                help="Upload one or more job descriptions",
                key="job_uploader"
            )

        all_files = []
        if resume_files:
            all_files.extend(self._process_uploaded_files(resume_files, "resume"))
        if job_files:
            all_files.extend(self._process_uploaded_files(job_files, "job_description"))
        
        return all_files
    
    def _process_uploaded_files(self, files, doc_type: str):
        """Process uploaded files with validation"""
        processed_files = []
        
        for file in files:
            validation_result = self._validate_file(file)
            
            if validation_result['is_valid']:
                processed_file = {
                    'name': file.name,
                    'type': doc_type,
                    'size': file.size,
                    'content': file.getvalue(),
                    'format': file.name.split('.')[-1].lower(),
                    'upload_time': datetime.now()
                }
                processed_files.append(processed_file)
            else:
                st.error(f"❌ {file.name}: {validation_result['error']}")
        
        return processed_files
    
    def _validate_file(self, file) -> Dict[str, any]:
        """Validate uploaded file"""
        if file.size > self.max_file_size:
            return {
                'is_valid': False, 
                'error': f'File size ({file.size/1024/1024:.1f}MB) exceeds 10MB limit'
            }
        
        if not file.name:
            return {'is_valid': False, 'error': 'Invalid file name'}
        
        file_extension = file.name.split('.')[-1].lower()
        if file_extension not in self.supported_formats:
            return {
                'is_valid': False, 
                'error': f'Unsupported format. Use: {", ".join(self.supported_formats)}'
            }
        
        if file.size == 0:
            return {'is_valid': False, 'error': 'File is empty'}
        
        return {'is_valid': True, 'error': None}


class TextExtractor:
    """Handles text extraction from different file formats"""
    
    def __init__(self):
        self.extraction_methods = {
            'pdf': self._extract_from_pdf,
            'docx': self._extract_from_docx,
            'txt': self._extract_from_txt
        }
        self.logger = self._setup_logger()
    
    def extract_text(self, file_info: Dict) -> Dict[str, any]:
        """Main text extraction method"""
        file_format = file_info['format']
        
        try:
            if file_format not in self.extraction_methods:
                raise ValueError(f"Unsupported format: {file_format}")
            
            extracted_text = self.extraction_methods[file_format](file_info['content'])
            
            if not extracted_text or len(extracted_text.strip()) < 10:
                raise ValueError("Extracted text is too short or empty")
            
            self.logger.info(f"Successfully extracted text from {file_info['name']}")
            
            return {
                'success': True,
                'text': extracted_text,
                'word_count': len(extracted_text.split()),
                'char_count': len(extracted_text),
                'extraction_method': file_format,
                'file_name': file_info['name']
            }
            
        except Exception as e:
            error_msg = f"Extraction failed for {file_info['name']}: {str(e)}"
            self.logger.error(error_msg)
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'word_count': 0,
                'char_count': 0,
                'file_name': file_info['name']
            }
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF files with fallback methods"""
        text = ""
        
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                except Exception as e:
                    self.logger.warning(f"Failed to extract page {page_num + 1}: {e}")
            
            if len(text.strip()) < 50:
                raise ValueError("PyPDF2 extraction yielded insufficient text")
            
        except Exception as e:
            try:
                import textract
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                    tmp.write(file_content)
                    tmp_path = tmp.name
                
                try:
                    text = textract.process(tmp_path).decode('utf-8')
                finally:
                    os.unlink(tmp_path)
            except ImportError:
                raise Exception("PDF extraction failed. Install textract: `pip install textract`")
            except Exception as fallback_error:
                raise Exception(f"All PDF extraction methods failed: {str(e)}, {str(fallback_error)}")
        
        return text.strip()
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(BytesIO(file_content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            text = "\n".join(text_parts)
            
            if not text.strip():
                raise ValueError("No text content found in DOCX file")
            
            return text
            
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def _extract_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT files with encoding detection"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                if len(text.strip()) > 0:
                    return text
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        try:
            return file_content.decode('utf-8', errors='replace')
        except Exception as e:
            raise Exception(f"TXT extraction failed: {str(e)}")
    
    def _setup_logger(self):
        """Setup logging for extraction process"""
        logger = logging.getLogger('TextExtractor')
        if not logger.handlers:
            logger.setLevel(logging.INFO)
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        return logger

class TextCleaner:
    """Handles text cleaning and preprocessing"""
    
    def __init__(self):
        self.cleaning_steps = [
            self._fix_hyphenated_words,
            self._remove_extra_whitespace,
            self._normalize_line_breaks,
            self._remove_special_characters,
            self._standardize_formatting,
            self._extract_sections
        ]
    
    def clean_text(self, raw_text: str, document_type: str = 'general') -> Dict[str, any]:
        """Main text cleaning pipeline"""
        if not raw_text or not raw_text.strip():
            return {
                'success': False,
                'error': 'Empty or invalid text provided',
                'cleaned_text': '',
                'cleaning_log': []
            }
        
        try:
            cleaned_text = raw_text
            cleaning_log = []
            
            for step in self.cleaning_steps:
                before_length = len(cleaned_text)
                cleaned_text = step(cleaned_text, document_type)
                after_length = len(cleaned_text)
                
                step_name = step.__name__.replace('_', ' ').title()
                cleaning_log.append({
                    'step': step_name,
                    'chars_before': before_length,
                    'chars_after': after_length,
                    'reduction': before_length - after_length
                })
            
            original_length = len(raw_text)
            final_length = len(cleaned_text)
            reduction_percentage = ((original_length - final_length) / original_length * 100) if original_length > 0 else 0
            
            return {
                'success': True,
                'cleaned_text': cleaned_text,
                'original_length': original_length,
                'final_length': final_length,
                'reduction_percentage': reduction_percentage,
                'cleaning_log': cleaning_log
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'cleaned_text': raw_text,
                'cleaning_log': []
            }
    
    def _fix_hyphenated_words(self, text: str, doc_type: str) -> str:
        """Find and rejoin words split by hyphens across lines"""
        return re.sub(r'(\w+)-\n(\w+)', r'\1\2', text)
    
    def _remove_extra_whitespace(self, text: str, doc_type: str) -> str:
        """Remove extra whitespace and normalize spacing"""
        text = re.sub(r' +', ' ', text)
        lines = [line.strip() for line in text.split('\n')]
        
        cleaned_lines = []
        prev_empty = False
        
        for line in lines:
            if line:
                cleaned_lines.append(line)
                prev_empty = False
            elif not prev_empty:
                cleaned_lines.append('')
                prev_empty = True
        
        return '\n'.join(cleaned_lines).strip()
    
    def _normalize_line_breaks(self, text: str, doc_type: str) -> str:
        """Normalize different line break formats"""
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text
    
    def _remove_special_characters(self, text: str, doc_type: str) -> str:
        """Remove unwanted special characters while preserving structure"""
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
        text = re.sub(r'[•·▪▫■□◦‣⁃→‰]', '• ', text)
        text = re.sub(r'[""\'\'«»]', '"', text)
        text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'^--- Page \d+ ---$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\d+$', '', text, flags=re.MULTILINE)
        return text
    
    def _standardize_formatting(self, text: str, doc_type: str) -> str:
        """Standardize common formatting patterns"""
        text = re.sub(
            r'\b(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})\b', 
            r'(\1) \2-\3', text
        )
        text = re.sub(r'\s*[|•]\s*', ' | ', text)
        return text
    
    def _extract_sections(self, text: str, doc_type: str) -> str:
        """Extract and organize document sections"""
        if doc_type == 'resume':
            return self._process_resume_sections(text)
        elif doc_type == 'job_description':
            return self._process_job_description_sections(text)
        return text

    def _process_resume_sections(self, text: str) -> str:
        """Process resume-specific sections."""
        section_patterns = [
            (r'\b(PROFESSIONAL\s+EXPERIENCE|WORK\s+EXPERIENCE|EXPERIENCE|EMPLOYMENT\s+HISTORY)\b', 'EXPERIENCE'),
            (r'\b(EDUCATION|EDUCATIONAL\s+BACKGROUND|ACADEMIC\s+BACKGROUND)\b', 'EDUCATION'),
            (r'\b(TECHNICAL\s+SKILLS|(?<!Soft\s)SKILLS|CORE\s+COMPETENCIES|COMPETENCIES)\b', 'SKILLS'),
            (r'\b(PROJECTS|PROJECT\s+EXPERIENCE|NOTABLE\s+PROJECTS)\b', 'PROJECTS'),
            (r'\b(CERTIFICATIONS|LICENSES|PROFESSIONAL\s+CERTIFICATIONS)\b', 'CERTIFICATIONS'),
            (r'\b(ACHIEVEMENTS|ACCOMPLISHMENTS|AWARDS)\b', 'ACHIEVEMENTS')
        ]
        
        sections = {}
        current_section = "UNCATEGORIZED"
        text_lines = text.split('\n')
        
        for line in text_lines:
            found_section = False
            
            for pattern, section_name in section_patterns:
                if re.search(pattern, line, flags=re.IGNORECASE):
                    current_section = section_name
                    if current_section not in sections:
                        sections[current_section] = []
                    found_section = True
                    break
            
            if not found_section:
                if current_section not in sections:
                    sections[current_section] = []
                sections[current_section].append(line)
        
        cleaned_text_parts = []
        
        for section_name in ["UNCATEGORIZED", "EDUCATION", "EXPERIENCE", "PROJECTS", "SKILLS", "CERTIFICATIONS", "ACHIEVEMENTS"]:
            if section_name in sections:
                if section_name not in ["UNCATEGORIZED"]:
                    cleaned_text_parts.append(f"\n\n=== {section_name} ===\n")
                cleaned_text_parts.append("\n".join(sections[section_name]))
        
        return "".join(cleaned_text_parts).strip()

    def _process_job_description_sections(self, text: str) -> str:
        """Process job description specific sections"""
        jd_patterns = [
            (r'\b(REQUIREMENTS|REQUIRED\s+QUALIFICATIONS|MINIMUM\s+REQUIREMENTS)\b', 'REQUIREMENTS'),
            (r'\b(RESPONSIBILITIES|JOB\s+DUTIES|KEY\s+RESPONSIBILITIES|DUTIES)\b', 'RESPONSIBILITIES'),
            (r'\b(PREFERRED|NICE\s+TO\s+HAVE|PREFERRED\s+QUALIFICATIONS)\b', 'PREFERRED'),
            (r'\b(BENEFITS|COMPENSATION|SALARY|PACKAGE)\b', 'BENEFITS'),
            (r'\b(ABOUT|COMPANY|ORGANIZATION|OVERVIEW)\b', 'ABOUT')
        ]

        sections = {}
        current_section = "UNCATEGORIZED"
        text_lines = text.split('\n')
        
        for line in text_lines:
            found_section = False
            for pattern, section_name in jd_patterns:
                if re.search(pattern, line, flags=re.IGNORECASE):
                    current_section = section_name
                    if current_section not in sections:
                        sections[current_section] = []
                    found_section = True
                    break
            
            if not found_section:
                if current_section not in sections:
                    sections[current_section] = []
                sections[current_section].append(line)
        
        cleaned_text_parts = []
        for section_name, content_lines in sections.items():
            if section_name != "UNCATEGORIZED":
                cleaned_text_parts.append(f"\n\n=== {section_name} ===\n")
            cleaned_text_parts.append("\n".join(content_lines))
        
        return "".join(cleaned_text_parts).strip()


class DocumentProcessor:
    """Main document processing pipeline coordinator"""
    
    def __init__(self):
        self.uploader = DocumentUploader()
        self.extractor = TextExtractor()
        self.cleaner = TextCleaner()
    
    def run_pipeline(self):
        """Execute the complete document processing pipeline"""
        if 'processed_docs' not in st.session_state:
            st.session_state.processed_docs = []

        st.title("🧠 AI Skill Gap Analyzer")
        st.markdown("A tool to analyze skills from resumes and job descriptions.")
        
        uploaded_files = self.uploader.create_upload_interface()
        
        st.markdown("---")
        
        if uploaded_files:
            if st.button("🚀 Process Documents", type="primary", use_container_width=True):
                self._process_all_documents(uploaded_files)
        else:
            st.info("⬆️ Upload files to begin processing.")
        
        if st.session_state.processed_docs:
            st.markdown("---")
            self._display_processing_results()
            self._create_download_interface()
    
    def _process_all_documents(self, uploaded_files: List[Dict]):
        """Process all uploaded documents"""
        progress_bar = st.progress(0, text="Initializing processing...")
        status_text = st.empty()
        
        processed_docs = []
        total_files = len(uploaded_files)
        
        for i, file_info in enumerate(uploaded_files):
            progress = (i + 1) / total_files
            progress_bar.progress(progress, text=f"Processing {file_info['name']} ({i+1}/{total_files})")
            
            result = self._process_single_document(file_info)
            processed_docs.append(result)
        
        st.session_state.processed_docs = processed_docs
        
        successful = sum(1 for doc in processed_docs if doc['success'])
        st.success(f"🎉 **Processing Complete!** {successful}/{total_files} documents processed successfully.")
    
    def _process_single_document(self, file_info: Dict) -> Dict:
        """Process a single document through the complete pipeline"""
        try:
            extraction_result = self.extractor.extract_text(file_info)
            
            if not extraction_result['success']:
                return {
                    'file_name': file_info['name'],
                    'document_type': file_info['type'],
                    'success': False,
                    'error': extraction_result['error'],
                    'stage': 'extraction',
                    'timestamp': datetime.now()
                }
            
            cleaning_result = self.cleaner.clean_text(
                extraction_result['text'], 
                file_info['type']
            )
            
            if not cleaning_result['success']:
                return {
                    'file_name': file_info['name'],
                    'document_type': file_info['type'],
                    'success': False,
                    'error': cleaning_result['error'],
                    'stage': 'cleaning',
                    'timestamp': datetime.now()
                }
            
            return {
                'file_name': file_info['name'],
                'document_type': file_info['type'],
                'success': True,
                'original_text': extraction_result['text'],
                'cleaned_text': cleaning_result['cleaned_text'],
                'extraction_stats': {
                    'word_count': extraction_result['word_count'],
                    'char_count': extraction_result['char_count'],
                    'extraction_method': extraction_result['extraction_method']
                },
                'cleaning_stats': {
                    'original_length': cleaning_result['original_length'],
                    'final_length': cleaning_result['final_length'],
                    'reduction_percentage': cleaning_result['reduction_percentage']
                },
                'processing_log': cleaning_result['cleaning_log'],
                'timestamp': datetime.now()
            }
            
        except Exception as e:
            return {
                'file_name': file_info['name'],
                'document_type': file_info['type'],
                'success': False,
                'error': str(e),
                'stage': 'general_processing',
                'timestamp': datetime.now()
            }
    
    def _display_processing_results(self):
        """Display comprehensive processing results"""
        st.header("📊 Processing Results")
        
        processed_docs = st.session_state.processed_docs
        success_docs = [doc for doc in processed_docs if doc['success']]
        failed_docs = [doc for doc in processed_docs if not doc['success']]
        
        tab1, tab2, tab3 = st.tabs([
            f"📈 Summary", 
            f"✅ Successful ({len(success_docs)})", 
            f"❌ Failed ({len(failed_docs)})"])
        
        with tab1:
            self._display_summary(success_docs, failed_docs)
            
        with tab2:
            if not success_docs:
                st.info("No documents were processed successfully.")
            for i, doc in enumerate(success_docs):
                self._display_successful_document(doc, i)
                
        with tab3:
            if not failed_docs:
                st.info("No documents failed to process.")
            for doc in failed_docs:
                self._display_failed_document(doc)

    def _display_summary(self, success_docs, failed_docs):
        """Display a summary of the processing results"""
        total = len(success_docs) + len(failed_docs)
        success_rate = (len(success_docs) / total * 100) if total > 0 else 0
        
        st.subheader("Overall Statistics")
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("📄 Total Documents", total)
        col2.metric("✅ Successful", len(success_docs))
        col3.metric("❌ Failed", len(failed_docs))
        col4.metric("🎯 Success Rate", f"{success_rate:.1f}%")

    def _display_successful_document(self, doc: Dict, index: int):
        """Display details for a successfully processed document"""
        with st.container():
            st.subheader(f"📄 {doc['file_name']} ({doc['document_type']})")
            
            col1, col2 = st.columns(2)
            with col1:
                st.text_area("Raw Text", doc['original_text'], height=300, key=f"raw_{index}", disabled=True)
            
            with col2:
                st.text_area("Cleaned Text", doc['cleaned_text'], height=300, key=f"cleaned_{index}", disabled=True)
            
            with st.expander("🔍 View Details"):
                st.subheader("Cleaning Steps")
                log_df = pd.DataFrame(doc['processing_log'])
                st.dataframe(log_df, use_container_width=True)

                st.subheader("Document Statistics")
                stats_col1, stats_col2, stats_col3, stats_col4 = st.columns(4)
                stats_col1.metric("Words", f"{doc['extraction_stats']['word_count']:,}")
                stats_col2.metric("Characters", f"{doc['extraction_stats']['char_count']:,}")
                stats_col3.metric("Method", doc['extraction_stats']['extraction_method'].upper())
                stats_col4.metric("Reduction", f"{doc['cleaning_stats']['reduction_percentage']:.1f}%")
            st.markdown("---")

    def _display_failed_document(self, doc: Dict):
        """Display details for a failed document"""
        with st.expander(f"❌ {doc['file_name']} - Failed at {doc['stage']}"):
            st.error(f"**Error:** {doc['error']}")
            st.write(f"**Document Type:** {doc['document_type']}")
            st.write(f"**Processing Stage:** {doc['stage']}")
            st.write(f"**Timestamp:** {doc['timestamp'].strftime('%Y-%m-%d %H:%M:%S')}")
    
    def _create_download_interface(self):
        """Create download interface with separate buttons for resumes and JDs."""
        successful_docs = [doc for doc in st.session_state.processed_docs if doc['success']]
    
        if not successful_docs:
            return
    
        st.header("⬇️ Download Processed Data")
        st.markdown("---") # Add a separator

        # --- NEW: Filter documents by type ---
        resume_docs = [doc for doc in successful_docs if doc['document_type'] == 'resume']
        jd_docs = [doc for doc in successful_docs if doc['document_type'] == 'job_description']

    # --- Section for Resumes ---
        if resume_docs:
            st.subheader(f"📄 Download Resumes ({len(resume_docs)})")
            col1, col2, col3 = st.columns(3)
        
            with col1:
                csv_data = self._create_csv_export(resume_docs)
                st.download_button(
                    label="⬇️ Download Resumes (CSV)",
                    data=csv_data,
                    file_name=f"processed_resumes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv",
                    use_container_width=True,
                    key ="download_resume_csv" # Add a unique key
                )
        
            with col2:
                json_data = self._create_json_export(resume_docs)
                st.download_button(
                    label="⬇️ Download Resumes (JSON)",
                    data=json_data,
                    file_name=f"processed_resumes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json",
                    use_container_width=True,
                    key="download_resume_json" # Add a unique key
                )
            
            with col3:
                text_data = self._create_text_export(resume_docs)
                st.download_button(
                    label="⬇️ Download Resumes (TXT)",
                    data=text_data,
                    file_name=f"processed_resumes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    use_container_width=True,
                    key="download_resume_txt" # Add a unique key
                )
            st.markdown("---") # Add a separator

    # --- Section for Job Descriptions ---
        if jd_docs:
            st.subheader(f"💼 Download Job Descriptions ({len(jd_docs)})")
            col1, col2, col3 = st.columns(3)
        
            with col1:
                csv_data = self._create_csv_export(jd_docs)
                st.download_button(
                    label="⬇️ Download JDs (CSV)",
                    data=csv_data,
                    file_name=f"processed_jds_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv",
                    use_container_width=True,
                    key="download_jd_csv" # Add a unique key
                )
        
            with col2:
                json_data = self._create_json_export(jd_docs)
                st.download_button(
                    label="⬇️ Download JDs (JSON)",
                    data=json_data,
                    file_name=f"processed_jds_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json",
                    use_container_width=True,
                    key="download_jd_json" # Add a unique key
                )

            with col3:
                text_data = self._create_text_export(jd_docs)
                st.download_button(
                    label="⬇️ Download JDs (TXT)",
                    data=text_data,
                    file_name=f"processed_jds_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    use_container_width=True,
                    key="download_jd_txt" # Add a unique key
                )

    def _create_csv_export(self, documents: List[Dict]) -> str:
            """Create CSV export of processed documents"""
            export_data = []
        
            for doc in documents:
                export_data.append({
                    'filename': doc['file_name'],
                    'document_type': doc['document_type'],
                    'word_count': doc['extraction_stats']['word_count'],
                    'char_count': doc['extraction_stats']['char_count'],
                    'original_length': doc['cleaning_stats']['original_length'],
                    'final_length': doc['cleaning_stats']['final_length'],
                    'reduction_percentage': doc['cleaning_stats']['reduction_percentage'],
                    'cleaned_text': doc['cleaned_text'],
                    'processing_timestamp': doc['timestamp'].strftime('%Y-%m-%d %H:%M:%S')
                })
        
            df = pd.DataFrame(export_data)
            return df.to_csv(index=False)
    
    def _create_json_export(self, documents: List[Dict]) -> str:
        """Create JSON export of processed documents"""
        export_data = []
        
        for doc in documents:
            clean_doc = {
                'filename': doc['file_name'],
                'document_type': doc['document_type'],
                'cleaned_text': doc['cleaned_text'],
                'extraction_stats': doc['extraction_stats'],
                'cleaning_stats': doc['cleaning_stats'],
                'processing_timestamp': doc['timestamp'].isoformat()
            }
            export_data.append(clean_doc)
        
        return json.dumps(export_data, indent=2)
    

    def _create_text_export(self, documents: List[Dict]) -> str:
        """Create a single text file export of all processed documents."""
        report_parts = []
        
        for doc in documents:
            report_parts.append("=" * 80)
            report_parts.append(f"FILE: {doc['file_name']}")
            report_parts.append(f"TYPE: {doc['document_type']}")
            report_parts.append("=" * 80)
            report_parts.append("\n")
            report_parts.append(doc['cleaned_text'])
            report_parts.append("\n\n")
            
        return "\n".join(report_parts)
def main():
    """Main application entry point"""
    try:
        processor = DocumentProcessor()
        
        with st.sidebar:
            st.header("🌐 Website Details")
            st.markdown("""
            This AI Skill Gap Analyzer supports:
            - **File Formats:** PDF, DOCX, TXT
            - **Max File Size:** 10MB per file
            
            **Processing Workflow:**
            1. Upload & Validate Files
            2. Extract Text
            3. Clean & Preprocess Text
            4. Display Results & Download
            """)
            st.markdown("---")
            
            st.header("📋 Status")
            if 'processed_docs' in st.session_state:
                total = len(st.session_state.processed_docs)
                successful = sum(1 for doc in st.session_state.processed_docs if doc['success'])
                st.success(f"Processed: {successful}/{total} documents")
            else:
                st.info("No documents processed yet")

        processor.run_pipeline()
        
    except Exception as e:
        st.error(f"Application error: {str(e)}")
        st.error("Please refresh the page and try again.")

if __name__ == "__main__":
    # Run the main function when the script is executed.
    main()