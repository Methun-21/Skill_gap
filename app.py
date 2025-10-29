import streamlit as st
import pandas as pd
import numpy as np
import PyPDF2
import docx
import re
import os
import logging
import tempfile
from typing import Dict, List, Set, Tuple, Optional
from io import BytesIO, StringIO
from datetime import datetime
import urllib.parse
import json
import random
from collections import Counter, defaultdict
from dataclasses import dataclass

# NLP and ML Imports
import spacy
from spacy.training import Example
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity

# Visualization Imports
import plotly.graph_objects as go
import plotly.express as px

# --- PAGE CONFIGURATION & STYLING ---
st.set_page_config(
    page_title="AI Skill Gap Analyzer - Final Application",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    :root {
        --primary-color: #3b82f6;
        --secondary-color: #8b5cf6;
        --accent-color: #10b981;
        --text-color: #e5e7eb;
        --bg-color: #111827;
        --card-bg-color: #1f2937;
        --border-color: #374151;
    }
    .stApp { background-color: var(--bg-color); color: var(--text-color); }
    h1 {
        color: #f9fafb;
        background: -webkit-linear-gradient(45deg, var(--primary-color), var(--secondary-color));
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    h2, h3 { color: #d1d5db; }
    .stButton > button {
        border-radius: 8px;
        background-image: linear-gradient(to right, var(--primary-color), var(--secondary-color));
        color: white; border: none; padding: 10px 24px; transition: all 0.3s ease;
        box-shadow: 0 4px 15px 0 rgba(139, 92, 246, 0.3);
    }
    .stButton > button:hover { transform: translateY(-2px); box-shadow: 0 6px 20px 0 rgba(139, 92, 246, 0.5); }
    button[data-baseweb="tab"][aria-selected="true"] {
        background-image: linear-gradient(to right, var(--primary-color), var(--secondary-color));
        color: white; border-radius: 8px 8px 0 0;
    }
    [data-testid="stExpander"], [data-testid="stMetric"], .stDataFrame, .custom-metric {
        background-color: var(--card-bg-color);
        border: 1px solid var(--border-color);
        border-radius: 10px;
        padding: 1rem;
        height: 100%;
    }
    .skill-tag { display: inline-block; padding: 6px 12px; margin: 4px; border-radius: 16px; font-weight: 500; border: 1px solid transparent; }
    .tech-skill { background-color: rgba(59, 130, 246, 0.2); border-color: #3b82f6; color: #93c5fd; }
    .soft-skill { background-color: rgba(168, 85, 247, 0.2); border-color: #a855f7; color: #d8b4fe; }

    /* --- NEW STYLES TO ADD --- */
    .metric-label {
        font-size: 1rem;
        color: var(--text-color);
        opacity: 0.8;
    }
    .metric-value-large {
        font-size: 2.25rem;
        line-height: 2.5rem;
        font-weight: 600;
        letter-spacing: -0.025em;
    }
    .green-text { color: #28a745 !important; }
    .yellow-text { color: #ffc107 !important; }
    .red-text { color: #dc3545 !important; }

</style>
""", unsafe_allow_html=True)

# ==============================================================================
# VERBATIM CLASS DEFINITIONS FROM ALL PROVIDED FILES
# ==============================================================================

# --- Classes from `complete_pipeline.py` ---

class DocumentUploader:
    """Handles file upload functionality with validation"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt']
        self.max_file_size = 10 * 1024 * 1024  # 10MB in bytes
    
    def create_upload_interface(self):
        """Create the main upload interface"""
        st.header("Upload Documents 📄")
        
        st.info("Upload your resumes and job descriptions here. The system will process the text and get it ready for analysis. Supported formats: PDF, DOCX, TXT.")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("📄 Resumes")
            resume_files = st.file_uploader(
                "Choose Resume files",
                type=self.supported_formats,
                accept_multiple_files=True,
                help="Upload one or more resume files",
                key="resume_uploader"
            )
            
        with col2:
            st.subheader("💼 Job Descriptions")
            job_files = st.file_uploader(
                "Choose Job Description files",
                type=self.supported_formats,
                accept_multiple_files=True,
                help="Upload one or more job descriptions",
                key="job_uploader"
            )

        all_files = []
        if resume_files:
            all_files.extend(self._process_uploaded_files(resume_files, "resume"))
        if job_files:
            all_files.extend(self._process_uploaded_files(job_files, "job_description"))
        
        return all_files
    
    def _process_uploaded_files(self, files, doc_type: str):
        """Process uploaded files with validation"""
        processed_files = []
        
        for file in files:
            validation_result = self._validate_file(file)
            
            if validation_result['is_valid']:
                processed_file = {
                    'name': file.name,
                    'type': doc_type,
                    'size': file.size,
                    'content': file.getvalue(),
                    'format': file.name.split('.')[-1].lower(),
                    'upload_time': datetime.now()
                }
                processed_files.append(processed_file)
            else:
                st.error(f"❌ {file.name}: {validation_result['error']}")
        
        return processed_files
    
    def _validate_file(self, file) -> Dict[str, any]:
        """Validate uploaded file"""
        if file.size > self.max_file_size:
            return {
                'is_valid': False, 
                'error': f'File size ({file.size/1024/1024:.1f}MB) exceeds 10MB limit'
            }
        
        if not file.name:
            return {'is_valid': False, 'error': 'Invalid file name'}
        
        file_extension = file.name.split('.')[-1].lower()
        if file_extension not in self.supported_formats:
            return {
                'is_valid': False, 
                'error': f'Unsupported format. Use: {", ".join(self.supported_formats)}'
            }
        
        if file.size == 0:
            return {'is_valid': False, 'error': 'File is empty'}
        
        return {'is_valid': True, 'error': None}

class TextExtractor:
    """Handles text extraction from different file formats"""
    
    def __init__(self):
        self.extraction_methods = {
            'pdf': self._extract_from_pdf,
            'docx': self._extract_from_docx,
            'txt': self._extract_from_txt
        }
        self.logger = self._setup_logger()
    
    def extract_text(self, file_info: Dict) -> Dict[str, any]:
        """Main text extraction method"""
        file_format = file_info['format']
        
        try:
            if file_format not in self.extraction_methods:
                raise ValueError(f"Unsupported format: {file_format}")
            
            extracted_text = self.extraction_methods[file_format](file_info['content'])
            
            if not extracted_text or len(extracted_text.strip()) < 10:
                raise ValueError("Extracted text is too short or empty")
            
            self.logger.info(f"Successfully extracted text from {file_info['name']}")
            
            return {
                'success': True,
                'text': extracted_text,
                'word_count': len(extracted_text.split()),
                'char_count': len(extracted_text),
                'extraction_method': file_format,
                'file_name': file_info['name']
            }
            
        except Exception as e:
            error_msg = f"Extraction failed for {file_info['name']}: {str(e)}"
            self.logger.error(error_msg)
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'word_count': 0,
                'char_count': 0,
                'file_name': file_info['name']
            }
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF files with fallback methods"""
        text = ""
        
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                except Exception as e:
                    self.logger.warning(f"Failed to extract page {page_num + 1}: {e}")
            
            if len(text.strip()) < 50:
                raise ValueError("PyPDF2 extraction yielded insufficient text")
            
        except Exception as e:
            try:
                import textract
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                    tmp.write(file_content)
                    tmp_path = tmp.name
                
                try:
                    text = textract.process(tmp_path).decode('utf-8')
                finally:
                    os.unlink(tmp_path)
            except ImportError:
                raise Exception("PDF extraction failed. Install textract: `pip install textract`")
            except Exception as fallback_error:
                raise Exception(f"All PDF extraction methods failed: {str(e)}, {str(fallback_error)}")
        
        return text.strip()
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(BytesIO(file_content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            text = "\n".join(text_parts)
            
            if not text.strip():
                raise ValueError("No text content found in DOCX file")
            
            return text
            
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def _extract_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT files with encoding detection"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                if len(text.strip()) > 0:
                    return text
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        try:
            return file_content.decode('utf-8', errors='replace')
        except Exception as e:
            raise Exception(f"TXT extraction failed: {str(e)}")
    
    def _setup_logger(self):
        """Setup logging for extraction process"""
        logger = logging.getLogger('TextExtractor')
        if not logger.handlers:
            logger.setLevel(logging.INFO)
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        return logger

class TextCleaner:
    """Handles text cleaning and preprocessing"""
    
    def __init__(self):
        self.cleaning_steps = [
            self._fix_hyphenated_words,
            self._remove_extra_whitespace,
            self._normalize_line_breaks,
            self._remove_special_characters,
            self._standardize_formatting,
            self._extract_sections
        ]
    
    def clean_text(self, raw_text: str, document_type: str = 'general') -> Dict[str, any]:
        """Main text cleaning pipeline"""
        if not raw_text or not raw_text.strip():
            return {
                'success': False,
                'error': 'Empty or invalid text provided',
                'cleaned_text': '',
                'cleaning_log': []
            }
        
        try:
            cleaned_text = raw_text
            cleaning_log = []
            
            for step in self.cleaning_steps:
                before_length = len(cleaned_text)
                cleaned_text = step(cleaned_text, document_type)
                after_length = len(cleaned_text)
                
                step_name = step.__name__.replace('_', ' ').title()
                cleaning_log.append({
                    'step': step_name,
                    'chars_before': before_length,
                    'chars_after': after_length,
                    'reduction': before_length - after_length
                })
            
            original_length = len(raw_text)
            final_length = len(cleaned_text)
            reduction_percentage = ((original_length - final_length) / original_length * 100) if original_length > 0 else 0
            
            return {
                'success': True,
                'cleaned_text': cleaned_text,
                'original_length': original_length,
                'final_length': final_length,
                'reduction_percentage': reduction_percentage,
                'cleaning_log': cleaning_log
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'cleaned_text': raw_text,
                'cleaning_log': []
            }
    
    def _fix_hyphenated_words(self, text: str, doc_type: str) -> str:
        """Find and rejoin words split by hyphens across lines"""
        return re.sub(r'(\w+)-\n(\w+)', r'\1\2', text)
    
    def _remove_extra_whitespace(self, text: str, doc_type: str) -> str:
        """Remove extra whitespace and normalize spacing"""
        text = re.sub(r' +', ' ', text)
        lines = [line.strip() for line in text.split('\n')]
        
        cleaned_lines = []
        prev_empty = False
        
        for line in lines:
            if line:
                cleaned_lines.append(line)
                prev_empty = False
            elif not prev_empty:
                cleaned_lines.append('')
                prev_empty = True
        
        return '\n'.join(cleaned_lines).strip()
    
    def _normalize_line_breaks(self, text: str, doc_type: str) -> str:
        """Normalize different line break formats"""
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text
    
    def _remove_special_characters(self, text: str, doc_type: str) -> str:
        """Remove unwanted special characters while preserving structure"""
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
        text = re.sub(r'[•·▪▫■□◦‣⁃→‰]', '• ', text)
        text = re.sub(r'[""\'\'«»]', '"', text)
        text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'^--- Page \d+ ---$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\d+$', '', text, flags=re.MULTILINE)
        return text
    
    def _standardize_formatting(self, text: str, doc_type: str) -> str:
        """Standardize common formatting patterns"""
        text = re.sub(
            r'\b(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})\b', 
            r'(\1) \2-\3', text
        )
        text = re.sub(r'\s*[|•]\s*', ' | ', text)
        return text
    
    def _extract_sections(self, text: str, doc_type: str) -> str:
        """Extract and organize document sections"""
        if doc_type == 'resume':
            return self._process_resume_sections(text)
        elif doc_type == 'job_description':
            return self._process_job_description_sections(text)
        return text

    def _process_resume_sections(self, text: str) -> str:
        """Process resume-specific sections."""
        section_patterns = [
            (r'\b(PROFESSIONAL\s+EXPERIENCE|WORK\s+EXPERIENCE|EXPERIENCE|EMPLOYMENT\s+HISTORY)\b', 'EXPERIENCE'),
            (r'\b(EDUCATION|EDUCATIONAL\s+BACKGROUND|ACADEMIC\s+BACKGROUND)\b', 'EDUCATION'),
            (r'\b(TECHNICAL\s+SKILLS|(?<!Soft\s)SKILLS|CORE\s+COMPETENCIES|COMPETENCIES)\b', 'SKILLS'),
            (r'\b(PROJECTS|PROJECT\s+EXPERIENCE|NOTABLE\s+PROJECTS)\b', 'PROJECTS'),
            (r'\b(CERTIFICATIONS|LICENSES|PROFESSIONAL\s+CERTIFICATIONS)\b', 'CERTIFICATIONS'),
            (r'\b(ACHIEVEMENTS|ACCOMPLISHMENTS|AWARDS)\b', 'ACHIEVEMENTS')
        ]
        
        sections = {}
        current_section = "UNCATEGORIZED"
        text_lines = text.split('\n')
        
        for line in text_lines:
            found_section = False
            
            for pattern, section_name in section_patterns:
                if re.search(pattern, line, flags=re.IGNORECASE):
                    current_section = section_name
                    if current_section not in sections:
                        sections[current_section] = []
                    found_section = True
                    break
            
            if not found_section:
                if current_section not in sections:
                    sections[current_section] = []
                sections[current_section].append(line)
        
        cleaned_text_parts = []
        
        for section_name in ["UNCATEGORIZED", "EDUCATION", "EXPERIENCE", "PROJECTS", "SKILLS", "CERTIFICATIONS", "ACHIEVEMENTS"]:
            if section_name in sections:
                if section_name not in ["UNCATEGORIZED"]:
                    cleaned_text_parts.append(f"\n\n=== {section_name} ===\n")
                cleaned_text_parts.append("\n".join(sections[section_name]))
        
        return "".join(cleaned_text_parts).strip()

    def _process_job_description_sections(self, text: str) -> str:
        """Process job description specific sections"""
        jd_patterns = [
            (r'\b(REQUIREMENTS|REQUIRED\s+QUALIFICATIONS|MINIMUM\s+REQUIREMENTS)\b', 'REQUIREMENTS'),
            (r'\b(RESPONSIBILITIES|JOB\s+DUTIES|KEY\s+RESPONSIBILITIES|DUTIES)\b', 'RESPONSIBILITIES'),
            (r'\b(PREFERRED|NICE\s+TO\s+HAVE|PREFERRED\s+QUALIFICATIONS)\b', 'PREFERRED'),
            (r'\b(BENEFITS|COMPENSATION|SALARY|PACKAGE)\b', 'BENEFITS'),
            (r'\b(ABOUT|COMPANY|ORGANIZATION|OVERVIEW)\b', 'ABOUT')
        ]

        sections = {}
        current_section = "UNCATEGORIZED"
        text_lines = text.split('\n')
        
        for line in text_lines:
            found_section = False
            for pattern, section_name in jd_patterns:
                if re.search(pattern, line, flags=re.IGNORECASE):
                    current_section = section_name
                    if current_section not in sections:
                        sections[current_section] = []
                    found_section = True
                    break
            
            if not found_section:
                if current_section not in sections:
                    sections[current_section] = []
                sections[current_section].append(line)
        
        cleaned_text_parts = []
        for section_name, content_lines in sections.items():
            if section_name != "UNCATEGORIZED":
                cleaned_text_parts.append(f"\n\n=== {section_name} ===\n")
            cleaned_text_parts.append("\n".join(content_lines))
        
        return "".join(cleaned_text_parts).strip()

class SkillDatabase:
    """Comprehensive skill database with categorization"""
    
    def __init__(self):
        self.skills = self._initialize_skill_database()
        self.abbreviations = self._initialize_abbreviations()
        self.skill_patterns = self._initialize_skill_patterns()
        self.non_skill_words = {
            'role', 'responsibilities', 'company', 'experience', 'summary',
            'education', 'requirements', 'skills', 'preferred', 'duties',
            'about', 'team', 'projects', 'work', 'contact', 'email', 'phone'
        }
    
    def _initialize_skill_database(self) -> Dict[str, List[str]]:
        """Initialize comprehensive skill database"""
        return {
            'programming_languages': [
                'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', 'C',
                'Ruby', 'PHP', 'Swift', 'Kotlin', 'Go', 'Rust', 'Scala', 'R',
                'MATLAB', 'Perl', 'Dart', 'Shell', 'Bash', 'PowerShell'
            ],
            'web_frameworks': [
                'React', 'Angular', 'Vue.js', 'Node.js', 'Express.js', 'Django',
                'Flask', 'FastAPI', 'Spring Boot', 'Spring', 'ASP.NET', '.NET Core',
                'Ruby on Rails', 'Laravel', 'Next.js', 'Nuxt.js', 'Svelte',
                'jQuery', 'Bootstrap', 'Tailwind CSS', 'Material-UI'
            ],
            'databases': [
                'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Cassandra', 'Oracle',
                'SQL Server', 'SQLite', 'MariaDB', 'DynamoDB', 'Elasticsearch',
                'Firebase', 'Neo4j', 'Snowflake', 'BigQuery'
            ],
            'ml_ai': [
                'Machine Learning', 'Deep Learning', 'Neural Networks',
                'Natural Language Processing', 'NLP', 'Computer Vision',
                'Reinforcement Learning', 'Transfer Learning',
                'Feature Engineering', 'MLOps', 'Generative AI',
                'Large Language Models', 'LLM', 'CNN', 'RNN', 'LSTM',
                'Transformer', 'BERT', 'GPT'
            ],
            'ml_frameworks': [
                'TensorFlow', 'PyTorch', 'Keras', 'Scikit-learn', 'XGBoost',
                'Pandas', 'NumPy', 'SciPy', 'Matplotlib', 'Seaborn',
                'Plotly', 'NLTK', 'spaCy', 'Hugging Face', 'OpenCV'
            ],
            'cloud_platforms': [
                'AWS', 'Amazon Web Services', 'Azure', 'Microsoft Azure',
                'Google Cloud Platform', 'GCP', 'Heroku', 'DigitalOcean'
            ],
            'devops_tools': [
                'Docker', 'Kubernetes', 'Jenkins', 'GitLab CI', 'GitHub Actions',
                'CircleCI', 'Ansible', 'Terraform', 'Prometheus', 'Grafana',
                'ELK Stack', 'Datadog'
            ],
            'version_control': [
                'Git', 'GitHub', 'GitLab', 'Bitbucket', 'SVN'
            ],
            'testing': [
                'Jest', 'Mocha', 'Pytest', 'JUnit', 'Selenium', 'Cypress',
                'Postman', 'JMeter'
            ],
            'soft_skills': [
                'Leadership', 'Team Management', 'Communication',
                'Problem Solving', 'Critical Thinking', 'Analytical Skills',
                'Project Management', 'Collaboration', 'Teamwork',
                'Adaptability', 'Creativity', 'Time Management'
            ]
        }
    
    def _initialize_abbreviations(self) -> Dict[str, str]:
        """Initialize common abbreviations"""
        return {
            'ML': 'Machine Learning', 'DL': 'Deep Learning',
            'AI': 'Artificial Intelligence', 'NLP': 'Natural Language Processing',
            'CV': 'Computer Vision', 'NN': 'Neural Networks',
            'CNN': 'Convolutional Neural Networks', 'RNN': 'Recurrent Neural Networks',
            'K8s': 'Kubernetes', 'K8S': 'Kubernetes',
            'CI/CD': 'Continuous Integration/Continuous Deployment',
            'API': 'Application Programming Interface',
            'REST': 'Representational State Transfer',
            "OS": "Operating System",
            "DB": "Database",
            "JS": "JavaScript",
            "HR": "Human Resources",
            'SQL': 'Structured Query Language', 'OOP': 'Object-Oriented Programming',
            'TDD': 'Test-Driven Development', 'AWS': 'Amazon Web Services',
            'GCP': 'Google Cloud Platform'
        }
    
    def _initialize_skill_patterns(self) -> List[str]:
        """
        Initialize regex patterns for skill detection.

        These patterns are designed to match common phrases in resumes and job descriptions
        that indicate skills, such as 'experience in Python', 'proficient with Django', or
        '5+ years of experience in Machine Learning'. The patterns aim to cover a wide range
        of phrasings to robustly extract both technical and soft skills from varied text sources.
        """
        return [
            r'experience (?:in|with) ([\w\s\+\#\.\-]+)',
            r'proficient (?:in|with|at) ([\w\s\+\#\.\-]+)',
            r'expertise (?:in|with) ([\w\s\+\#\.\-]+)',
            r'knowledge of ([\w\s\+\#\.\-]+)',
            r'skilled (?:in|at|with) ([\w\s\+\#\.\-]+)',
            r'familiar with ([\w\s\+\#\.\-]+)',
            r'(\d+)\+?\s*years? of (?:experience )?(?:in|with) ([\w\s\+\#\.\-]+)'
        ]
    
    def get_all_skills(self) -> List[str]:
        """Get flattened list of all skills"""
        all_skills = []
        for skills in self.skills.values():
            all_skills.extend(skills)
        return all_skills
    
    def get_category_for_skill(self, skill: str) -> Optional[str]:
        """Find which category a skill belongs to"""
        skill_lower = skill.lower()
        for category, skills in self.skills.items():
            if any(s.lower() == skill_lower for s in skills):
                return category
        return 'other'
    
class TextPreprocessor:
    """Advanced text preprocessing for skill extraction"""
    
    def __init__(self):
        self.nlp = self._load_spacy_model()
        self._customize_stop_words()
    
    def _load_spacy_model(self):
        """Load spaCy model with error handling"""
        try:
            return spacy.load("en_core_web_sm")
        except OSError:
            st.warning("⚠️ Installing spaCy model...")
            import subprocess
            subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])
            return spacy.load("en_core_web_sm")
    
    def _customize_stop_words(self):
        """Customize stop words for skill extraction"""
        programming_langs = {'c', 'r', 'go', 'd', 'f'}
        for lang in programming_langs:
            self.nlp.Defaults.stop_words.discard(lang)
    
    def preprocess(self, text: str) -> Dict:
        """Complete preprocessing pipeline"""
        if not text or not text.strip():
            return {'success': False, 'error': 'Empty text'}
        
        try:
            doc = self.nlp(text)
            
            return {
                'success': True,
                'doc': doc,
                'noun_chunks': [chunk.text for chunk in doc.noun_chunks],
                'entities': [(ent.text, ent.label_) for ent in doc.ents],
                'sentences': [sent.text for sent in doc.sents]
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}

class SkillExtractor:
    """Main skill extraction engine"""
    
    def __init__(self):
        self.skill_db = SkillDatabase()
        self.preprocessor = TextPreprocessor()
        self.logger = self._setup_logger()
    
    def extract_skills(self, text: str, document_type: str = 'resume') -> Dict:
        """Extract skills using multiple methods"""
        try:
            preprocess_result = self.preprocessor.preprocess(text)
            if not preprocess_result['success']:
                return {'success': False, 'error': preprocess_result['error']}
            
            doc = preprocess_result['doc']
            
            # Multiple extraction methods
            keyword_skills = self._extract_by_keywords(text)
            pos_skills = self._extract_by_pos_patterns(doc)
            context_skills = self._extract_by_context(text)
            ner_skills = self._extract_by_ner(preprocess_result['entities'])
            chunk_skills = self._extract_from_noun_chunks(preprocess_result['noun_chunks'])
            
            # Combine and deduplicate
            all_skills = self._combine_and_deduplicate([
                keyword_skills, pos_skills, context_skills, ner_skills, chunk_skills
            ])
            
            # Normalize (expand abbreviations)
            normalized_skills = self._normalize_skills(all_skills)
            
            # Categorize skills
            categorized_skills = self._categorize_skills(normalized_skills)
            
            # Calculate confidence scores
            skill_confidence = self._calculate_confidence(
                normalized_skills,
                [keyword_skills, pos_skills, context_skills, ner_skills, chunk_skills]
            )
            
            return {
                'success': True,
                'all_skills': normalized_skills,
                'categorized_skills': categorized_skills,
                'skill_confidence': skill_confidence,
                'extraction_methods': {
                    'keyword_matching': len(keyword_skills),
                    'pos_patterns': len(pos_skills),
                    'context_based': len(context_skills),
                    'ner': len(ner_skills),
                    'noun_chunks': len(chunk_skills)
                },
                'statistics': {
                    'total_skills': len(normalized_skills),
                    'technical_skills': sum(len(skills) for cat, skills in categorized_skills.items() if cat != 'soft_skills'),
                    'soft_skills': len(categorized_skills.get('soft_skills', []))
                }
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _extract_by_keywords(self, text: str) -> Set[str]:
        """Extract skills by keyword matching"""
        found_skills = set()
        text_lower = text.lower()
        
        for skill in self.skill_db.get_all_skills():
            pattern = r'\b' + re.escape(skill.lower()) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.add(skill)
        
        return found_skills
    
    def _extract_by_pos_patterns(self, doc) -> Set[str]:
        """Extract skills using POS patterns"""
        found_skills = set()
        tokens = list(doc)
        
        # ADJ + NOUN patterns
        for i in range(len(tokens) - 1):
            if tokens[i].pos_ == 'ADJ' and tokens[i+1].pos_ in ['NOUN', 'PROPN']:
                pattern = f"{tokens[i].text} {tokens[i+1].text}"
                if self._is_valid_skill(pattern):
                    found_skills.add(pattern)
        
        # Proper nouns
        for token in doc:
            if token.pos_ == 'PROPN' and self._is_valid_skill(token.text):
                found_skills.add(token.text)
        
        return found_skills
    
    def _extract_by_context(self, text: str) -> Set[str]:
        """Extract skills based on context patterns"""
        found_skills = set()
        
        for pattern in self.skill_db.skill_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                if len(match.groups()) >= 1:
                    skill_text = match.group(len(match.groups()))
                    skills = self._clean_and_split_skills(skill_text)
                    for skill in skills:
                        if self._is_valid_skill(skill):
                            found_skills.add(skill)
        
        return found_skills
    
    def _extract_by_ner(self, entities: List[Tuple[str, str]]) -> Set[str]:
        """Extract skills from named entities"""
        found_skills = set()
        relevant_labels = ['ORG', 'PRODUCT', 'GPE']
        
        for entity_text, label in entities:
            if label in relevant_labels and self._is_valid_skill(entity_text):
                found_skills.add(entity_text)
        
        return found_skills
    
    def _extract_from_noun_chunks(self, noun_chunks: List[str]) -> Set[str]:
        """Extract skills from noun chunks"""
        found_skills = set()
        
        for chunk in noun_chunks:
            chunk_clean = chunk.strip()
            if self._is_valid_skill(chunk_clean):
                found_skills.add(chunk_clean)
        
        return found_skills
    
    def _is_valid_skill(self, text: str) -> bool:
        """
        Validate if the given text is a valid skill using a strict, exact-match method.
        """
        text_clean = text.strip()
        text_lower = text_clean.lower()

        # 1. Check against the list of non-skill words (blacklist).
        if text_lower in self.skill_db.non_skill_words:
            return False

        # 2. Perform a basic check for length. Most single-letter words are not skills,
        # unless they are explicitly in the database (like 'C' or 'R').
        if not text_clean or len(text_clean) < 1:
            return False
    
        # 3. The primary validation: Check for an exact, case-insensitive match
        # in the comprehensive skill database.
        all_skills_lower = [s.lower() for s in self.skill_db.get_all_skills()]
        if text_lower in all_skills_lower:
            return True

        # If the text is not found as an exact match in the skill database,
        # it is not considered a valid skill.
        return False
    def _clean_and_split_skills(self, text: str) -> List[str]:
        """Clean and split comma-separated skills"""
        skills = re.split(r'[,;|/&]|\band\b', text)
        cleaned_skills = []
        
        for skill in skills:
            skill_clean = skill.strip()
            skill_clean = re.sub(r'\b(etc|and more)\b', '', skill_clean, flags=re.IGNORECASE).strip()
            if skill_clean and len(skill_clean) > 1:
                cleaned_skills.append(skill_clean)
        
        return cleaned_skills
    
    def _combine_and_deduplicate(self, skill_sets: List[Set[str]]) -> List[str]:
        """Combine and remove duplicates"""
        combined = set()
        for skill_set in skill_sets:
            combined.update(skill_set)
        
        unique_skills = {}
        for skill in combined:
            skill_lower = skill.lower()
            if skill_lower not in unique_skills:
                unique_skills[skill_lower] = skill
        
        return sorted(unique_skills.values())
    
    def _normalize_skills(self, skills: List[str]) -> List[str]:
        """Normalize skill names (expand abbreviations)"""
        normalized = []
        
        for skill in skills:
            if skill.upper() in self.skill_db.abbreviations:
                normalized.append(self.skill_db.abbreviations[skill.upper()])
            else:
                normalized.append(skill)
        
        return sorted(set(normalized))
    
    def _categorize_skills(self, skills: List[str]) -> Dict[str, List[str]]:
        """Categorize skills"""
        categorized = defaultdict(list)
        
        for skill in skills:
            category = self.skill_db.get_category_for_skill(skill)
            categorized[category].append(skill)
        
        for category in categorized:
            categorized[category] = sorted(categorized[category])
        
        return dict(categorized)
    
    def _calculate_confidence(self, skills: List[str], method_results: List[Set[str]]) -> Dict[str, float]:
        """Calculate confidence score for each skill"""
        confidence_scores = {}
        
        for skill in skills:
            detection_count = sum(
                1 for method_set in method_results 
                if skill in method_set or skill.lower() in {s.lower() for s in method_set}
            )
            confidence = detection_count / len(method_results)
            confidence_scores[skill] = round(confidence, 2)
        
        return confidence_scores
    
    def _setup_logger(self):
        """Setup logging"""
        logger = logging.getLogger('SkillExtractor')
        if not logger.handlers:
            logger.setLevel(logging.INFO)
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        return logger
# --- Classes from `milestone3.py` ---

@dataclass
class SkillMatch:
    """Individual skill match details"""
    jd_skill: str
    resume_skill: str
    similarity: float
    category: str
    confidence_level: str
    priority: str = "MEDIUM"
    
    def to_dict(self) -> Dict:
        return {
            'jd_skill': self.jd_skill,
            'resume_skill': self.resume_skill,
            'similarity': self.similarity,
            'category': self.category,
            'confidence_level': self.confidence_level,
            'priority': self.priority
        }

@dataclass
class GapAnalysisResult:
    """Complete gap analysis results"""
    matched_skills: List[SkillMatch]
    partial_matches: List[SkillMatch]
    missing_skills: List[SkillMatch]
    overall_score: float
    category_scores: Dict[str, float]
    similarity_matrix: np.ndarray
    resume_skills: List[str]
    jd_skills: List[str]
    
    def get_statistics(self) -> Dict:
        total = len(self.jd_skills)
        return {
            'total_required_skills': total,
            'matched_count': len(self.matched_skills),
            'partial_count': len(self.partial_matches),
            'missing_count': len(self.missing_skills),
            'match_percentage': (len(self.matched_skills) / total * 100) if total > 0 else 0,
            'overall_score': self.overall_score * 100
        }

class SentenceBERTEncoder:
    """Handles BERT embedding generation using Sentence-BERT"""
    
    def __init__(self, model_name: str = 'all-MiniLM-L6-v2'):
        """
        Initialize Sentence-BERT model
        """
        self.model_name = model_name
        self.logger = self._setup_logger()
        self.embedding_cache = {}
        
        try:
            self.logger.info(f"Loading model: {model_name}")
            self.model = SentenceTransformer(model_name)
            self.embedding_dimension = self.model.get_sentence_embedding_dimension()
            self.logger.info(f"Model loaded successfully. Embedding dimension: {self.embedding_dimension}")
        except Exception as e:
            self.logger.error(f"Failed to load model: {e}")
            raise
    
    def encode_skills(self, skills: List[str], use_cache: bool = True, 
                     show_progress: bool = False) -> np.ndarray:
        if not skills:
            raise ValueError("Skills list cannot be empty")
        
        # Check cache
        if use_cache:
            cached_embeddings = []
            uncached_skills = []
            uncached_indices = []
            
            for i, skill in enumerate(skills):
                if skill in self.embedding_cache:
                    cached_embeddings.append(self.embedding_cache[skill])
                else:
                    uncached_skills.append(skill)
                    uncached_indices.append(i)
            
            # Encode uncached skills
            if uncached_skills:
                new_embeddings = self.model.encode(
                    uncached_skills, 
                    show_progress_bar=show_progress,
                    batch_size=32
                )
                
                # Update cache
                for skill, embedding in zip(uncached_skills, new_embeddings):
                    self.embedding_cache[skill] = embedding
                
                # Combine cached and new
                all_embeddings = [None] * len(skills)
                cached_idx = 0
                uncached_idx = 0
                
                for i in range(len(skills)):
                    if i in uncached_indices:
                        all_embeddings[i] = new_embeddings[uncached_idx]
                        uncached_idx += 1
                    else:
                        all_embeddings[i] = cached_embeddings[cached_idx]
                        cached_idx += 1
                
                return np.array(all_embeddings)
            else:
                return np.array(cached_embeddings)
        else:
            # Encode without cache
            embeddings = self.model.encode(
                skills, 
                show_progress_bar=show_progress,
                batch_size=32
            )
            return embeddings
    
    def get_embedding_for_skill(self, skill: str) -> np.ndarray:
        """Get embedding for a single skill"""
        if skill in self.embedding_cache:
            return self.embedding_cache[skill]
        
        embedding = self.model.encode([skill])[0]
        self.embedding_cache[skill] = embedding
        return embedding
    
    def clear_cache(self):
        """Clear embedding cache"""
        self.embedding_cache.clear()
        self.logger.info("Embedding cache cleared")
    
    def _setup_logger(self) -> logging.Logger:
        """Setup logging"""
        logger = logging.getLogger('BERTEncoder')
        if not logger.handlers:
            logger.setLevel(logging.INFO)
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        return logger

class SimilarityCalculator:
    """Compute similarity scores between skills"""
    
    def __init__(self):
        self.logger = self._setup_logger()
    
    def compute_cosine_similarity(self, embedding1: np.ndarray, 
                                  embedding2: np.ndarray) -> float:
        # Reshape if needed
        if embedding1.ndim == 1:
            embedding1 = embedding1.reshape(1, -1)
        if embedding2.ndim == 1:
            embedding2 = embedding2.reshape(1, -1)
        
        similarity = cosine_similarity(embedding1, embedding2)[0][0]
        return float(similarity)
    
    def compute_similarity_matrix(self, resume_embeddings: np.ndarray,
                                  jd_embeddings: np.ndarray) -> np.ndarray:
        self.logger.info(f"Computing similarity matrix: {resume_embeddings.shape} x {jd_embeddings.shape}")
        if resume_embeddings.shape[0] == 0 or jd_embeddings.shape[0] == 0:
            return np.array([[]]) # Return empty matrix if no skills
        similarity_matrix = cosine_similarity(resume_embeddings, jd_embeddings)
        self.logger.info(f"Similarity matrix computed: {similarity_matrix.shape}")
        return similarity_matrix
    
    def find_best_matches(self, similarity_matrix: np.ndarray, 
                         threshold: float = 0.5) -> List[Tuple[int, int, float]]:
        matches = []
        n_resume, n_jd = similarity_matrix.shape
        
        for jd_idx in range(n_jd):
            best_resume_idx = np.argmax(similarity_matrix[:, jd_idx])
            best_similarity = similarity_matrix[best_resume_idx, jd_idx]
            
            if best_similarity >= threshold:
                matches.append((best_resume_idx, jd_idx, best_similarity))
        
        return matches
    
    def _setup_logger(self) -> logging.Logger:
        """Setup logging"""
        logger = logging.getLogger('SimilarityCalculator')
        if not logger.handlers:
            logger.setLevel(logging.INFO)
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        return logger

class SkillGapAnalyzer:
    """Main skill gap analysis engine"""
    
    def __init__(self, encoder: SentenceBERTEncoder, calculator: SimilarityCalculator,
                 strong_threshold: float = 0.80, partial_threshold: float = 0.50):
        self.encoder = encoder
        self.calculator = calculator
        self.strong_threshold = strong_threshold
        self.partial_threshold = partial_threshold
        self.logger = self._setup_logger()
    
    def analyze(self, resume_skills: List[str], jd_skills: List[str],
               skill_categories: Optional[Dict[str, str]] = None) -> GapAnalysisResult:
        self.logger.info(f"Starting gap analysis: {len(resume_skills)} resume skills vs {len(jd_skills)} JD skills")
        
        # Validate inputs
        if not resume_skills or not jd_skills:
            raise ValueError("Both resume_skills and jd_skills must be non-empty")
        
        # Step 1: Generate embeddings
        self.logger.info("Step 1: Generating BERT embeddings...")
        resume_embeddings = self.encoder.encode_skills(resume_skills, show_progress=True)
        jd_embeddings = self.encoder.encode_skills(jd_skills, show_progress=True)
        
        # Step 2: Compute similarity matrix
        self.logger.info("Step 2: Computing similarity matrix...")
        similarity_matrix = self.calculator.compute_similarity_matrix(
            resume_embeddings, 
            jd_embeddings
        )
        
        # Step 3: Classify matches
        self.logger.info("Step 3: Classifying skill matches...")
        matched_skills = []
        partial_matches = []
        missing_skills = []
        
        for jd_idx, jd_skill in enumerate(jd_skills):
            # Find best matching resume skill
            best_resume_idx = np.argmax(similarity_matrix[:, jd_idx])
            best_similarity = float(similarity_matrix[best_resume_idx, jd_idx])
            resume_skill = resume_skills[best_resume_idx]
            
            # Get category
            category = skill_categories.get(jd_skill, 'other') if skill_categories else 'other'
            
            # Classify based on similarity
            if best_similarity >= self.strong_threshold:
                match = SkillMatch(
                    jd_skill=jd_skill,
                    resume_skill=resume_skill,
                    similarity=best_similarity,
                    category='STRONG_MATCH',
                    confidence_level='HIGH',
                    priority='LOW'
                )
                matched_skills.append(match)
                
            elif best_similarity >= self.partial_threshold:
                match = SkillMatch(
                    jd_skill=jd_skill,
                    resume_skill=resume_skill,
                    similarity=best_similarity,
                    category='PARTIAL_MATCH',
                    confidence_level='MEDIUM',
                    priority='MEDIUM'
                )
                partial_matches.append(match)
                
            else:
                match = SkillMatch(
                    jd_skill=jd_skill,
                    resume_skill=resume_skill,
                    similarity=best_similarity,
                    category='MISSING',
                    confidence_level='LOW',
                    priority='HIGH'
                )
                missing_skills.append(match)
        
        # Step 4: Calculate overall score
        overall_score = self._calculate_overall_score(similarity_matrix)
        
        # Step 5: Calculate category scores
        category_scores = self._calculate_category_scores(
            matched_skills, partial_matches, missing_skills
        )
        
        self.logger.info(f"Analysis complete: {len(matched_skills)} matched, "
                        f"{len(partial_matches)} partial, {len(missing_skills)} missing")
        
        return GapAnalysisResult(
            matched_skills=matched_skills,
            partial_matches=partial_matches,
            missing_skills=missing_skills,
            overall_score=overall_score,
            category_scores=category_scores,
            similarity_matrix=similarity_matrix,
            resume_skills=resume_skills,
            jd_skills=jd_skills
        )
    
    def _calculate_overall_score(self, similarity_matrix: np.ndarray) -> float:
        """Calculate overall match score"""
        # Take maximum similarity for each JD skill
        max_similarities = similarity_matrix.max(axis=0)
        # Average of all maximum similarities
        overall_score = float(np.mean(max_similarities))
        return overall_score
    
    def _calculate_category_scores(self, matched: List[SkillMatch],
                                   partial: List[SkillMatch],
                                   missing: List[SkillMatch]) -> Dict[str, float]:
        """Calculate scores by category"""
        category_scores = {}
        
        all_skills = matched + partial + missing
        categories = set(skill.category for skill in all_skills)
        
        for category in categories:
            cat_skills = [s for s in all_skills if s.category == category]
            if cat_skills:
                avg_similarity = np.mean([s.similarity for s in cat_skills])
                category_scores[category] = float(avg_similarity)
        
        return category_scores
    
    def _setup_logger(self) -> logging.Logger:
        """Setup logging"""
        logger = logging.getLogger('SkillGapAnalyzer')
        if not logger.handlers:
            logger.setLevel(logging.INFO)
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        return logger

class SkillRanker:
    """Rank skills by importance and priority"""
    
    def __init__(self):
        self.logger = self._setup_logger()
    
    def rank_by_importance(self, skills: List[SkillMatch], 
                          importance_weights: Optional[Dict[str, float]] = None) -> List[SkillMatch]:
        """
        Rank skills by importance
        
        Args:
            skills: List of SkillMatch objects
            importance_weights: Optional weights for different factors
            
        Returns:
            Ranked list of skills
        """
        if not importance_weights:
            importance_weights = {
                'similarity': 0.4,
                'category': 0.3,
                'priority': 0.3
            }
        
        def calculate_importance_score(skill: SkillMatch) -> float:
            # Similarity score
            sim_score = skill.similarity
            
            # Category score (missing skills are more important to address)
            if skill.category == 'MISSING':
                cat_score = 1.0
            elif skill.category == 'PARTIAL_MATCH':
                cat_score = 0.6
            else:
                cat_score = 0.2
            
            # Priority score
            priority_map = {'HIGH': 1.0, 'MEDIUM': 0.6, 'LOW': 0.3}
            pri_score = priority_map.get(skill.priority, 0.5)
            
            # Weighted combination
            importance = (
                importance_weights['similarity'] * sim_score +
                importance_weights['category'] * cat_score +
                importance_weights['priority'] * pri_score
            )
            
            return importance
        
        # Sort by importance (descending)
        ranked_skills = sorted(skills, key=calculate_importance_score, reverse=True)
        
        self.logger.info(f"Ranked {len(skills)} skills by importance")
        return ranked_skills
    
    def categorize_by_urgency(self, missing_skills: List[SkillMatch]) -> Dict[str, List[SkillMatch]]:
        """
        Categorize missing skills by urgency
        
        Returns:
            Dictionary with 'critical', 'important', and 'beneficial' keys
        """
        categorized = {
            'critical': [],
            'important': [],
            'beneficial': []
        }
        
        for skill in missing_skills:
            if skill.priority == 'HIGH' or skill.similarity < 0.3:
                categorized['critical'].append(skill)
            elif skill.priority == 'MEDIUM' or skill.similarity < 0.4:
                categorized['important'].append(skill)
            else:
                categorized['beneficial'].append(skill)
        
        return categorized
    
    def _setup_logger(self) -> logging.Logger:
        """Setup logging"""
        logger = logging.getLogger('SkillRanker')
        if not logger.handlers:
            logger.setLevel(logging.INFO)
        return logger

class GapVisualizer:
    """Create visualizations for gap analysis"""
    
    @staticmethod
    def create_similarity_heatmap(similarity_matrix: np.ndarray,
                                     resume_skills: List[str],
                                     jd_skills: List[str]) -> go.Figure:
        """Create a clean, interactive similarity heatmap."""

        # Limit display to avoid overcrowding
        max_display = 20
        display_resume = resume_skills[:max_display]
        display_jd = jd_skills[:max_display]
        display_matrix = similarity_matrix[:max_display, :max_display]

    # Create hover text to show the skill pair and score
        hover_text = []
        for r_skill in display_resume:
            row = []
            for j_skill in display_jd:
                row.append(f"Resume: {r_skill}<br>JD: {j_skill}")
            hover_text.append(row)

        fig = go.Figure(data=go.Heatmap(
            z=display_matrix,
            x=display_jd,
            y=display_resume,
            colorscale='RdYlGn',
            zmid=0.5,
            texttemplate='%{z:.2f}',
            textfont={"size": 10},
            hoverinfo='z+text',
            hovertext=hover_text,
            colorbar=dict(
                title=dict(
                    text="Similarity",
                    side="right"
                )
            )
        ))

        # 3. Refine the layout for better readability
        fig.update_layout(
            title_text=f"Skill Similarity Heatmap (Top {min(max_display, len(resume_skills))} Skills)",
            title_x=0.5,
            xaxis_title="Job Description Skills",
            yaxis_title="Resume Skills",
            height=700,
            width=900,
            xaxis={'side': 'bottom', 'tickangle': -45},
            yaxis={'autorange': 'reversed'},
            margin=dict(l=120, r=40, b=120, t=80) # Adjust margins
        )

        return fig
    
    @staticmethod
    def create_match_distribution_pie(analysis_result: GapAnalysisResult) -> go.Figure:
        """Create pie chart for match distribution"""
        
        stats = analysis_result.get_statistics()
        
        labels = ['Strong Matches', 'Partial Matches', 'Missing Skills']
        values = [stats['matched_count'], stats['partial_count'], stats['missing_count']]
        colors = ['#28a745', '#ffc107', '#dc3545']
        
        fig = go.Figure(data=[go.Pie(
            labels=labels,
            values=values,
            marker=dict(colors=colors),
            hole=0.3,
            textposition='auto',
            textinfo='label+percent+value'
        )])
        
        fig.update_layout(
            title="Skill Match Distribution",
            height=500,
            showlegend=True
        )
        
        return fig
    
    @staticmethod
    def create_skill_comparison_bar(analysis_result: GapAnalysisResult, top_n: int = 15) -> go.Figure:
        """Create bar chart comparing skill similarities"""
        
        all_matches = (analysis_result.matched_skills + 
                      analysis_result.partial_matches + 
                      analysis_result.missing_skills)
        
        # Sort by similarity
        all_matches_sorted = sorted(all_matches, key=lambda x: x.similarity, reverse=True)[:top_n]
        
        skills = [m.jd_skill for m in all_matches_sorted]
        similarities = [m.similarity * 100 for m in all_matches_sorted]
        colors_map = {'STRONG_MATCH': '#28a745', 'PARTIAL_MATCH': '#ffc107', 'MISSING': '#dc3545'}
        colors = [colors_map[m.category] for m in all_matches_sorted]
        
        fig = go.Figure(data=[go.Bar(
            y=skills,
            x=similarities,
            orientation='h',
            marker=dict(color=colors),
            text=[f"{s:.1f}%" for s in similarities],
            textposition='auto'
        )])
        
        fig.update_layout(
            title=f"Top {top_n} Skills by Similarity Score",
            xaxis_title="Similarity Score (%)",
            yaxis_title="Skills",
            height=600,
            yaxis=dict(autorange="reversed"),
            showlegend=False
        )
        
        return fig
    
    @staticmethod
    def create_gap_priority_chart(missing_skills: List[SkillMatch]) -> go.Figure:
        """Create chart showing gap priorities"""
        
        if not missing_skills:
            # Return empty figure if no missing skills
            fig = go.Figure()
            fig.update_layout(
                title="No Missing Skills",
                annotations=[dict(
                    text="All required skills are matched!",
                    xref="paper", yref="paper",
                    x=0.5, y=0.5, showarrow=False,
                    font=dict(size=20)
                )]
            )
            return fig
        
        # Sort by similarity (lower similarity = higher priority)
        sorted_skills = sorted(missing_skills, key=lambda x: x.similarity)[:15]
        
        skills = [s.jd_skill for s in sorted_skills]
        similarities = [s.similarity * 100 for s in sorted_skills]
        priorities = [s.priority for s in sorted_skills]
        
        priority_colors = {
            'HIGH': '#dc3545',
            'MEDIUM': '#ffc107',
            'LOW': '#28a745'
        }
        colors = [priority_colors[p] for p in priorities]
        
        fig = go.Figure(data=[go.Bar(
            y=skills,
            x=similarities,
            orientation='h',
            marker=dict(color=colors),
            text=[f"{s:.1f}% - {p}" for s, p in zip(similarities, priorities)],
            textposition='auto'
        )])
        
        fig.update_layout(
            title="Missing Skills by Priority",
            xaxis_title="Current Similarity (%)",
            yaxis_title="Skills",
            height=500,
            yaxis=dict(autorange="reversed")
        )
        
        return fig
    
    @staticmethod
    def create_overall_score_gauge(overall_score: float) -> go.Figure:
        """Create gauge chart for overall match score"""
        
        score_percentage = overall_score * 100
        
        fig = go.Figure(go.Indicator(
            mode="gauge+number+delta",
            value=score_percentage,
            domain={'x': [0, 1], 'y': [0, 1]},
            title={'text': "Overall Match Score", 'font': {'size': 24}},
            delta={'reference': 70, 'increasing': {'color': "green"}},
            gauge={
                'axis': {'range': [None, 100], 'tickwidth': 1, 'tickcolor': "darkblue"},
                'bar': {'color': "darkblue"},
                'bgcolor': "white",
                'borderwidth': 2,
                'bordercolor': "gray",
                'steps': [
                    {'range': [0, 40], 'color': '#ffcccc'},
                    {'range': [40, 70], 'color': '#ffffcc'},
                    {'range': [70, 100], 'color': '#ccffcc'}
                ],
                'threshold': {
                    'line': {'color': "red", 'width': 4},
                    'thickness': 0.75,
                    'value': 70
                }
            }
        ))
        
        fig.update_layout(
            height=400,
            margin=dict(l=20, r=20, t=50, b=20)
        )
        
        return fig

class ReportGenerator:
    """Generate comprehensive reports"""
    
    def __init__(self):
        self.timestamp = datetime.now()
    
    def generate_text_report(self, analysis_result: GapAnalysisResult) -> str:
        """Generate detailed text report"""
        
        stats = analysis_result.get_statistics()
        
        report_lines = []
        report_lines.append("=" * 80)
        report_lines.append("SKILL GAP ANALYSIS REPORT")
        report_lines.append("=" * 80)
        report_lines.append(f"\nGenerated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        report_lines.append("")
        
        # Executive Summary
        report_lines.append("-" * 80)
        report_lines.append("EXECUTIVE SUMMARY")
        report_lines.append("-" * 80)
        report_lines.append(f"Overall Match Score: {stats['overall_score']:.1f}%")
        report_lines.append(f"Total Required Skills: {stats['total_required_skills']}")
        report_lines.append(f"Matched Skills: {stats['matched_count']} ({stats['match_percentage']:.1f}%)")
        report_lines.append(f"Partial Matches: {stats['partial_count']}")
        report_lines.append(f"Missing Skills: {stats['missing_count']}")
        report_lines.append("")
        
        # Strong Matches
        if analysis_result.matched_skills:
            report_lines.append("-" * 80)
            report_lines.append("✓ STRONG MATCHES (Similarity ≥ 80%)")
            report_lines.append("-" * 80)
            for match in analysis_result.matched_skills:
                report_lines.append(f"  • {match.jd_skill}")
                report_lines.append(f"    Resume: {match.resume_skill}")
                report_lines.append(f"    Similarity: {match.similarity*100:.1f}%")
                report_lines.append("")
        
        # Partial Matches
        if analysis_result.partial_matches:
            report_lines.append("-" * 80)
            report_lines.append("⚠ PARTIAL MATCHES (Similarity 50-80%)")
            report_lines.append("-" * 80)
            for match in analysis_result.partial_matches:
                report_lines.append(f"  • {match.jd_skill}")
                report_lines.append(f"    Closest: {match.resume_skill}")
                report_lines.append(f"    Similarity: {match.similarity*100:.1f}%")
                report_lines.append(f"    Recommendation: Strengthen knowledge in {match.jd_skill}")
                report_lines.append("")
        
        # Missing Skills
        if analysis_result.missing_skills:
            report_lines.append("-" * 80)
            report_lines.append("✗ CRITICAL GAPS (Similarity < 50%)")
            report_lines.append("-" * 80)
            for match in analysis_result.missing_skills:
                report_lines.append(f"  • {match.jd_skill} - {match.priority} PRIORITY")
                report_lines.append(f"    Current closest: {match.resume_skill} ({match.similarity*100:.1f}%)")
                report_lines.append(f"    Action: Acquire {match.jd_skill} through training/certification")
                report_lines.append("")
        
        report_lines.append("=" * 80)
        report_lines.append("END OF REPORT")
        report_lines.append("=" * 80)
        
        return "\n".join(report_lines)
    
    def generate_csv_report(self, analysis_result: GapAnalysisResult) -> str:
        """Generate CSV report"""
        
        data = []
        
        # Add all matches
        for match in analysis_result.matched_skills:
            data.append({
                'JD Skill': match.jd_skill,
                'Resume Skill': match.resume_skill,
                'Similarity (%)': f"{match.similarity*100:.2f}",
                'Category': match.category,
                'Priority': match.priority,
                'Status': 'Matched'
            })
        
        for match in analysis_result.partial_matches:
            data.append({
                'JD Skill': match.jd_skill,
                'Resume Skill': match.resume_skill,
                'Similarity (%)': f"{match.similarity*100:.2f}",
                'Category': match.category,
                'Priority': match.priority,
                'Status': 'Partial'
            })
        
        for match in analysis_result.missing_skills:
            data.append({
                'JD Skill': match.jd_skill,
                'Resume Skill': match.resume_skill,
                'Similarity (%)': f"{match.similarity*100:.2f}",
                'Category': match.category,
                'Priority': match.priority,
                'Status': 'Missing'
            })
        
        df = pd.DataFrame(data)
        return df.to_csv(index=False)
    
    def generate_json_report(self, analysis_result: GapAnalysisResult) -> str:
        """Generate JSON report"""
        
        stats = analysis_result.get_statistics()
        
        report_data = {
            'timestamp': self.timestamp.isoformat(),
            'statistics': stats,
            'matched_skills': [match.to_dict() for match in analysis_result.matched_skills],
            'partial_matches': [match.to_dict() for match in analysis_result.partial_matches],
            'missing_skills': [match.to_dict() for match in analysis_result.missing_skills],
            'category_scores': analysis_result.category_scores,
            'resume_skills': analysis_result.resume_skills,
            'jd_skills': analysis_result.jd_skills
        }
        
        return json.dumps(report_data, indent=2)

class LearningPathGenerator:
    """Generate personalized learning paths for skill gaps"""

    def __init__(self):
        self.resource_database = self._initialize_resources()

    def _initialize_resources(self) -> Dict:
        """Initialize a small, curated database of learning resources."""
        return {
            'Python': {
                'difficulty': 'Medium',
                'time_estimate': '4-8 weeks',
                'resources': [
                    '[Problem Solving using Python](https://infyspringboard.onwingspan.com/web/en/app/toc/lex_32448858012147495000/overview)',
                    '[Python for Data Science](https://infyspringboard.onwingspan.com/web/en/app/toc/lex_1910241031275886600/overview)'
                ]
            },
            'Machine Learning': {
                'difficulty': 'Hard',
                'time_estimate': '12-16 weeks',
                'prerequisites': ['Python', 'Statistics'],
                'resources': [
                    '[Machine Learning Introduction](https://infyspringboard.onwingspan.com/web/en/app/toc/lex_34451216570228390000/overview)'
                ]
            },
            'AWS': {
                'difficulty': 'Medium',
                'time_estimate': '8-12 weeks',
                'resources': [
                    '[AWS Cloud Practitioner Essentials](https://infyspringboard.onwingspan.com/web/en/app/toc/lex_33804818511475250000/overview)'
                ]
            }
        }

    def _generate_springboard_search_link(self, skill: str) -> str:
        """Generates a dynamic search link for a skill on Infosys Springboard."""
        # Corrected base URL and query parameter structure
        base_url = "https://infyspringboard.onwingspan.com/web/en/app/search"
        # URL-encode the skill to handle spaces and special characters like C++
        encoded_skill = urllib.parse.quote_plus(skill)
        # The new URL structure uses '?q=' for the query
        search_url = f"{base_url}?q={encoded_skill}"
        return f"[Search for '{skill}' courses on Infosys Springboard]({search_url})"

    def generate_path(self, missing_skills: List[SkillMatch], current_skills: List[str]) -> List[Dict]:
        """Generate learning path for missing skills, with fallback search links."""
        learning_plan = []
        # Sort by priority
        sorted_skills = sorted(missing_skills, key=lambda x: (0 if x.priority == 'HIGH' else 1 if x.priority == 'MEDIUM' else 2, x.similarity))

        for skill_match in sorted_skills:
            skill = skill_match.jd_skill
            plan_item = {
                'skill': skill,
                'current_similarity': skill_match.similarity,
                'priority': skill_match.priority,
                'difficulty': 'Unknown',
                'time_estimate': 'Varies',
                'resources': [],
                'prerequisites': [],
                'missing_prerequisites': []
            }

            # Check if we have curated info for this skill in our small database
            if skill in self.resource_database:
                resource_info = self.resource_database[skill]
                plan_item.update({
                    'difficulty': resource_info.get('difficulty', 'Unknown'),
                    'time_estimate': resource_info.get('time_estimate', 'Varies'),
                    'resources': resource_info.get('resources', []),
                    'prerequisites': resource_info.get('prerequisites', [])
                })
                # Always add the dynamic search link as an additional option
                plan_item['resources'].append(self._generate_springboard_search_link(skill))
            else:
                # If the skill is not in our database, provide the dynamic search link as the primary resource
                plan_item['resources'] = [self._generate_springboard_search_link(skill)]

            # Check for missing prerequisites
            missing_prereqs = []
            for prereq in plan_item['prerequisites']:
                if prereq.lower() not in [s.lower() for s in current_skills]:
                    missing_prereqs.append(prereq)
            plan_item['missing_prerequisites'] = missing_prereqs
            
            learning_plan.append(plan_item)
            
        return learning_plan

# ==============================================================================
# MAIN STREAMLIT APPLICATION
# ==============================================================================

@st.cache_resource
def load_spacy_model(model_name="en_core_web_sm"):
    try: return spacy.load(model_name)
    except OSError:
        st.error(f"SpaCy model not found. Please run: python -m spacy download {model_name}")
        st.stop()

@st.cache_resource
def load_bert_encoder(model_name='all-MiniLM-L6-v2'):
    return SentenceBERTEncoder(model_name)

class MainApp:
    def __init__(self):
        # Load heavy models once
        self.nlp_model = load_spacy_model()
        self.bert_encoder = load_bert_encoder()
        
        # Instantiate all component classes
        self.text_extractor = TextExtractor()
        self.text_cleaner = TextCleaner()
        self.skill_extractor = SkillExtractor()
        self.similarity_calculator = SimilarityCalculator()
        self.skill_ranker = SkillRanker()
        self.gap_visualizer = GapVisualizer()
        self.report_generator = ReportGenerator()
        self.learning_path_generator = LearningPathGenerator()

        # Initialize session state
        self.init_session_state()

    def init_session_state(self):
        defaults = {
            'processed_resumes': None, 'processed_jd': None,
            'extracted_resume_skills': None, 'extracted_jd_skills': None,
            'analysis_result': None, 'similarity_threshold': 0.65, 'strong_threshold': 0.85
        }
        for key, val in defaults.items():
            if key not in st.session_state: st.session_state[key] = val

    def run(self):
        st.title("🚀 AI Skill Gap Analyzer")
        self.render_sidebar()
        self.render_main_content()

    def render_sidebar(self):
        with st.sidebar:
            st.header("⚙️ Controls")
            st.subheader("1. Upload Documents")
            resume_files = st.file_uploader("Upload Resumes", type=['pdf', 'docx', 'txt'], accept_multiple_files=True)
            jd_file = st.file_uploader("Upload Job Description", type=['pdf', 'docx', 'txt'])

            st.subheader("2. Configure Analysis")
            st.session_state.strong_threshold = st.slider("Strong Match Threshold", 0.0, 1.0, 0.80, 0.05)
            st.session_state.similarity_threshold = st.slider("Partial Match Threshold", 0.0, 1.0, 0.60, 0.05)

            st.subheader("3. Actions")
            if st.button("▶️ Run Full Analysis", type="primary", use_container_width=True):
                self.run_full_pipeline(resume_files, jd_file)
            if st.button("🔄 Reset", use_container_width=True):
                self.init_session_state()
                st.rerun()

    def run_full_pipeline(self, resume_files, jd_file):
        if not resume_files or not jd_file:
            st.error("Please upload at least one resume and one job description.")
            return

        with st.status("Step 1: Parsing and Cleaning Documents...", expanded=True) as status:
            processed_resumes, processed_jd = {}, {}
            # Process Resumes
            for file in resume_files:
                file_info = {'name': file.name, 'content': file.getvalue(), 'format': file.name.split('.')[-1].lower()}
                extract_res = self.text_extractor.extract_text(file_info)
                if extract_res['success']:
                    clean_res = self.text_cleaner.clean_text(extract_res['text'], 'resume')
                    # Store all necessary stats for the new UI
                    processed_resumes[file.name] = {
                        'raw': extract_res['text'],
                        'cleaned': clean_res['cleaned_text'],
                        'status': '✅',
                        'word_count': extract_res.get('word_count', 0),
                        'char_count': clean_res.get('final_length', 0),
                        'reduction_percentage': clean_res.get('reduction_percentage', 0.0)
                    }
                else:
                    processed_resumes[file.name] = {'status': '❌'}
            st.session_state.processed_resumes = processed_resumes

            # Process JD
            file_info = {'name': jd_file.name, 'content': jd_file.getvalue(), 'format': jd_file.name.split('.')[-1].lower()}
            extract_res = self.text_extractor.extract_text(file_info)
            if extract_res['success']:
                clean_res = self.text_cleaner.clean_text(extract_res['text'], 'job_description')
                processed_jd[jd_file.name] = {
                    'raw': extract_res['text'],
                    'cleaned': clean_res['cleaned_text'],
                    'status': '✅',
                    'word_count': extract_res.get('word_count', 0),
                    'char_count': clean_res.get('final_length', 0),
                    'reduction_percentage': clean_res.get('reduction_percentage', 0.0)
                }
            else:
                processed_jd[jd_file.name] = {'status': '❌'}
            st.session_state.processed_jd = processed_jd
            status.update(label="✅ Documents Processed!", state="complete")

        with st.status("Step 2: Extracting Skills with NLP...", expanded=True) as status:
            resume_text = " ".join([d['cleaned'] for d in processed_resumes.values() if 'cleaned' in d])
            st.session_state.extracted_resume_skills = self.skill_extractor.extract_skills(resume_text)
            jd_text = list(processed_jd.values())[0]['cleaned']
            st.session_state.extracted_jd_skills = self.skill_extractor.extract_skills(jd_text)
            status.update(label="✅ Skills Extracted!", state="complete")

        with st.status("Step 3: Analyzing Skill Gaps with BERT...", expanded=True) as status:
            analyzer = SkillGapAnalyzer(self.bert_encoder, self.similarity_calculator,
                                        st.session_state.strong_threshold, st.session_state.similarity_threshold)
            all_resume_skills = st.session_state.extracted_resume_skills['all_skills']
            all_jd_skills = st.session_state.extracted_jd_skills['all_skills']
        
            # CORRECTED: Pass the categorized skills from the JD
            jd_categorized_skills = st.session_state.extracted_jd_skills['categorized_skills']
            skill_to_category_map = {skill: cat for cat, skills in jd_categorized_skills.items() for skill in skills}

            original_result = analyzer.analyze(all_resume_skills, all_jd_skills, skill_to_category_map)
            st.session_state.analysis_result = original_result
            st.session_state.original_analysis_result = original_result

            status.update(label="✅ Gap Analysis Complete!", state="complete")
    
        st.balloons()
        st.toast("Analysis complete! The dashboard is ready.", icon="🎉")
        st.success("Full analysis pipeline completed!")

    def render_main_content(self):
        if not st.session_state.analysis_result:
            st.info("👋 Welcome! Upload documents and click 'Run Full Analysis' to begin.")
            return

        tabs = st.tabs(["📊 Dashboard", "📄 Documents", "🛠️ Skills", "📈 Gap Analysis", "🎓 Learning Path", "📥 Export"])
        with tabs[0]: self.render_dashboard_tab()
        with tabs[1]: self.render_documents_tab()
        with tabs[2]: self.render_skills_tab()
        with tabs[3]: self.render_gap_analysis_tab()
        with tabs[4]: self.render_learning_path_tab()
        with tabs[5]: self.render_export_tab()

    def render_dashboard_tab(self):
        st.header("📊 Analysis Dashboard")
        res = st.session_state.analysis_result
        stats = res.get_statistics()

        # --- Top Metrics With Custom Colors ---
        col1, col2, col3, col4, col5 = st.columns(5)
    
        # Use standard metrics for non-status cards
        col1.metric("Overall Match Score", f"{stats['overall_score']:.1f}%")
        col2.metric("Total JD Skills", stats['total_required_skills'])

        # Use custom markdown with CSS classes for the colored metrics
        with col3:
            st.markdown(f"""
            <div class="custom-metric">
                <div class="metric-label">Strong Matches</div>
                <div class="metric-value-large green-text">{stats['matched_count']}</div>
            </div>
            """, unsafe_allow_html=True)

        with col4:
            st.markdown(f"""
            <div class="custom-metric">
                <div class="metric-label">Partial Matches</div>
                <div class="metric-value-large yellow-text">{stats['partial_count']}</div>
            </div>
            """, unsafe_allow_html=True)

        with col5:
            st.markdown(f"""
            <div class="custom-metric">
                <div class="metric-label">Missing Skills</div>
                <div class="metric-value-large red-text">{stats['missing_count']}</div>
            </div>
            """, unsafe_allow_html=True)
    
        st.markdown("---")
    
        # --- Main Visualizations (Gauge and Simulator) ---
        c1, c2 = st.columns([1, 1])
        with c1:
            st.plotly_chart(self.gap_visualizer.create_overall_score_gauge(res.overall_score), use_container_width=True)

        with c2:
            st.subheader("🧪 What-If Skill Simulator")
            st.info("Add a skill you plan to learn and see how it impacts your score.")
        
            if 'original_analysis_result' not in st.session_state:
                st.session_state.original_analysis_result = res

            new_skill = st.text_input("Enter a skill to simulate adding to your profile:", placeholder="e.g., Docker")

            sim_col1, sim_col2 = st.columns(2)
            with sim_col1:
                if st.button("📈 Simulate Score", use_container_width=True, type="primary"):
                    if new_skill:
                        with st.spinner("Re-calculating analysis..."):
                            # Get current skills and add the new one
                            current_resume_skills = st.session_state.extracted_resume_skills['all_skills']
                            updated_resume_skills = current_resume_skills + [new_skill] if new_skill not in current_resume_skills else current_resume_skills

                            # Re-run only the analysis part
                            analyzer = SkillGapAnalyzer(self.bert_encoder, self.similarity_calculator, 
                                                        st.session_state.strong_threshold, st.session_state.similarity_threshold)
                        
                            all_jd_skills = st.session_state.extracted_jd_skills['all_skills']
                        
                            # CRITICAL FIX: Create and pass the skill category map to the analyzer
                            jd_categorized = st.session_state.extracted_jd_skills['categorized_skills']
                            skill_to_category_map = {skill: cat for cat, skills in jd_categorized.items() for skill in skills}
                        
                            st.session_state.analysis_result = analyzer.analyze(updated_resume_skills, all_jd_skills, skill_to_category_map)
                            st.toast(f"Simulated adding '{new_skill}'! Check your new score.", icon="✨")
                            st.rerun()
                    else:
                        st.warning("Please enter a skill to simulate.")
        
            with sim_col2:
                if st.button("🔄 Reset Analysis", use_container_width=True):
                    st.session_state.analysis_result = st.session_state.original_analysis_result
                    st.toast("Analysis reset to the original state.", icon="🔄")
                    st.rerun()

    def render_documents_tab(self):
        st.header("📄 Document Viewer")
        if st.session_state.processed_resumes and st.session_state.processed_jd:
            col1, col2 = st.columns(2)
            with col1:
                st.subheader("Resumes")
                for name, data in st.session_state.processed_resumes.items():
                    st.markdown(f"**File:** `{name}`")
                
                    # Create the metric cards in a row
                    stat_cols = st.columns(3)
                    stat_cols[0].metric("Words", f"{data.get('word_count', 0):,}")
                    stat_cols[1].metric("Characters", f"{data.get('char_count', 0):,}")
                    stat_cols[2].metric("Reduction", f"{data.get('reduction_percentage', 0.0):.1f}%")

                    with st.expander("View Cleaned Text"):
                        st.text_area("Cleaned Text", data.get('cleaned', ''), height=300, key=f"resume_{name}", label_visibility="collapsed")
                    st.markdown("---")


            with col2:
                st.subheader("Job Description")
                name, data = list(st.session_state.processed_jd.items())[0]
                st.markdown(f"**File:** `{name}`")

                # Create the metric cards for the JD
                stat_cols = st.columns(3)
                stat_cols[0].metric("Words", f"{data.get('word_count', 0):,}")
                stat_cols[1].metric("Characters", f"{data.get('char_count', 0):,}")
                stat_cols[2].metric("Reduction", f"{data.get('reduction_percentage', 0.0):.1f}%")

                with st.expander("View Cleaned Text", expanded=True):
                    st.text_area("Cleaned Text", data.get('cleaned', ''), height=300, key=f"jd_{name}", label_visibility="collapsed")
    def render_skills_tab(self):
        st.header("🛠️ Extracted Skills")
        res, jd = st.session_state.extracted_resume_skills, st.session_state.extracted_jd_skills
        if not res or not jd: return

        col1, col2 = st.columns(2)
        with col1:
            st.subheader(f"Resume Skills ({res['statistics']['total_skills']})")
            for cat, skills in res['categorized_skills'].items():
                st.markdown(f"**{cat.replace('_', ' ').title()}**")
                html = "".join([f'<span class="skill-tag tech-skill">{s}</span>' for s in skills])
                st.markdown(html, unsafe_allow_html=True)
        with col2:
            st.subheader(f"JD Skills ({jd['statistics']['total_skills']})")
            for cat, skills in jd['categorized_skills'].items():
                st.markdown(f"**{cat.replace('_', ' ').title()}**")
                html = "".join([f'<span class="skill-tag soft-skill">{s}</span>' for s in skills])
                st.markdown(html, unsafe_allow_html=True)
                
    def render_gap_analysis_tab(self):
        st.header("📈 Detailed Gap Analysis")
        res = st.session_state.analysis_result
    
        # --- Visualization Section (No Changes) ---
        st.plotly_chart(self.gap_visualizer.create_match_distribution_pie(res), use_container_width=True)
        st.markdown("---")
    
        # --- Interactive Table Section ---
        st.subheader("Explore Skill Match Details")

        # Create the full DataFrame first
        df_data = []
        for m in res.matched_skills:
            df_data.append({'JD Skill': m.jd_skill, 'Resume Match': m.resume_skill, 'Similarity': m.similarity, 'Status': 'Strong'})
        for p in res.partial_matches:
            df_data.append({'JD Skill': p.jd_skill, 'Resume Match': p.resume_skill, 'Similarity': p.similarity, 'Status': 'Partial'})
        for n in res.missing_skills:
            df_data.append({'JD Skill': n.jd_skill, 'Resume Match': n.resume_skill, 'Similarity': n.similarity, 'Status': 'Missing'})
    
        if not df_data:
            st.warning("No skill match data to display.")
            return

        df = pd.DataFrame(df_data).sort_values('Similarity', ascending=False)
    
        # --- Interactive Filter Widgets ---
        col1, col2 = st.columns([2, 3])
        with col1:
            # Filter by Status
            status_options = df['Status'].unique()
            selected_statuses = st.multiselect(
                "Filter by Status:",
                options=status_options,
                default=status_options
            )
        with col2:
            # UPDATED: Filter by Similarity Score as a percentage
            similarity_range_percent = st.slider(
                "Filter by Similarity Score:",
                min_value=0,
                max_value=100,
                value=(0, 100),  # Default to the full 0-100 range
                step=1,
                format="%d%%"  # Display the value as an integer percentage
            )

        # Convert the selected percentage range back to the 0.0 to 1.0 float scale for filtering
        min_similarity = similarity_range_percent[0] / 100.0
        max_similarity = similarity_range_percent[1] / 100.0

        # --- Apply Filters to the DataFrame using the converted float values ---
        filtered_df = df[
            df['Status'].isin(selected_statuses) &
            (df['Similarity'] >= min_similarity) &
            (df['Similarity'] <= max_similarity)
        ]

        # Display the filtered and styled table
        st.dataframe(
            filtered_df.style.format({'Similarity': '{:.1%}'}), 
            use_container_width=True,
            height=500
        )
    
        st.markdown("---")
        st.subheader("Similarity Heatmap")
        st.plotly_chart(self.gap_visualizer.create_similarity_heatmap(res.similarity_matrix, res.resume_skills, res.jd_skills), use_container_width=True)

    def render_learning_path_tab(self):
        st.header("🎓 Personalized Learning Path")
        res = st.session_state.analysis_result
    
        # Combine partial and missing skills to identify all gaps
        gaps = res.partial_matches + res.missing_skills
    
        if not gaps:
            st.success("🎉 No significant skill gaps found! Your profile is a strong match.")
            return
    
        st.info("Here is a suggested learning plan to address the identified skill gaps, prioritized by importance.")
    
        # Generate the learning path
        learning_plan = self.learning_path_generator.generate_path(gaps, res.resume_skills)
    
        for i, item in enumerate(learning_plan):
            # Use an expander for each skill to keep the UI clean
            with st.expander(f"**{i+1}. Learn {item['skill']}** (Priority: {item['priority']})"):
            
                st.markdown(f"**Current Similarity:** `{item['current_similarity']*100:.1f}%`")
                st.markdown(f"**Estimated Difficulty:** `{item['difficulty']}`")
                st.markdown(f"**Estimated Time:** `{item['time_estimate']}`")
            
                # Display any missing prerequisites
                if item['missing_prerequisites']:
                    st.warning(f"**Prerequisites Needed:** {', '.join(item['missing_prerequisites'])}")
            
                st.markdown("**Recommended Resources:**")
                for r in item['resources']:
                    st.markdown(f"- {r}")

    def render_export_tab(self):
        st.header("📥 Export Reports")
        res = st.session_state.analysis_result
        st.info("Download the complete analysis in your preferred format.")

        # Generate all report formats
        csv_report = self.report_generator.generate_csv_report(res)
        txt_report = self.report_generator.generate_text_report(res)
        json_report = self.report_generator.generate_json_report(res)

        # Create columns for the download buttons
        col1, col2, col3 = st.columns(3)
        with col1:
            st.download_button(
                label="⬇️ Download as CSV",
                data=csv_report,
                file_name="skill_gap_report.csv",
                mime="text/csv",
                use_container_width=True
            )
        with col2:
            st.download_button(
                label="⬇️ Download as TXT",
                data=txt_report,
                file_name="skill_gap_report.txt",
                mime="text/plain",
                use_container_width=True
            )
        with col3:
            st.download_button(
                label="⬇️ Download as JSON",
                data=json_report,
                file_name="skill_gap_report.json",
                mime="application/json",
                use_container_width=True
            )

        st.subheader("Preview CSV Report")
        df = pd.read_csv(StringIO(csv_report))
        st.dataframe(df)

if __name__ == "__main__":
    app = MainApp()
    app.run()