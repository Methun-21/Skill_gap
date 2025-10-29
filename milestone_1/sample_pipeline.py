"""
Complete Document Processing Pipeline for AI Skill Gap Analyzer
Milestone 1: Data Ingestion and Parsing

Requirements:
pip install streamlit PyPDF2 python-docx textract pandas

Run with: streamlit run document_processor.py
"""

import streamlit as st
import PyPDF2
import docx
import pandas as pd
import re
import os
import logging
import tempfile
from typing import Dict, List, Optional, Tuple
from io import BytesIO, StringIO
from datetime import datetime
import json

# Configure page settings
st.set_page_config(
    page_title="AI Skill Gap Analyzer - Document Processor",
    page_icon="📄",
    layout="wide"
)

# Custom CSS for a cleaner, more spacious look


class DocumentUploader:
    """Handles file upload functionality with validation"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt']
        self.max_file_size = 10 * 1024 * 1024  # 10MB in bytes
    
    def create_upload_interface(self):
        """Create the main upload interface"""
        st.title("🔍 AI Skill Gap Analyzer - Document Processing")
        st.markdown("Upload resumes and job descriptions for skill gap analysis")
        
        # Create columns for different document types
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("📄 Resume Upload")
            resume_files = st.file_uploader(
                "Choose Resume files",
                type=self.supported_formats,
                accept_multiple_files=True,
                help="Upload resume in PDF, DOCX, or TXT format",
                key="resume_uploader"
            )
            
        with col2:
            st.subheader("💼 Job Description Upload")
            job_files = st.file_uploader(
                "Choose Job Description files",
                type=self.supported_formats,
                accept_multiple_files=True,
                help="Upload job description in PDF, DOCX, or TXT format",
                key="job_uploader"
            )

        
        # Process uploaded files
        all_files = []
        if resume_files:
            all_files.extend(self._process_uploaded_files(resume_files, "resume"))
        if job_files:
            all_files.extend(self._process_uploaded_files(job_files, "job_description"))
        
        return all_files
    
    def _process_uploaded_files(self, files, doc_type: str):
        """Process uploaded files with validation"""
        processed_files = []
        
        for file in files:
            validation_result = self._validate_file(file)
            
            if validation_result['is_valid']:
                processed_file = {
                    'name': file.name,
                    'type': doc_type,
                    'size': file.size,
                    'content': file.getvalue(),
                    'format': file.name.split('.')[-1].lower(),
                    'upload_time': datetime.now()
                }
                processed_files.append(processed_file)
            else:
                st.error(f"❌ {file.name}: {validation_result['error']}")
        
        return processed_files
    
    def _validate_file(self, file) -> Dict[str, any]:
        """Validate uploaded file"""
        # Check file size
        if file.size > self.max_file_size:
            return {
                'is_valid': False, 
                'error': f'File size ({file.size/1024/1024:.1f}MB) exceeds 10MB limit'
            }
        
        # Check file format
        if not file.name:
            return {'is_valid': False, 'error': 'Invalid file name'}
        
        file_extension = file.name.split('.')[-1].lower()
        if file_extension not in self.supported_formats:
            return {
                'is_valid': False, 
                'error': f'Unsupported format. Use: {", ".join(self.supported_formats)}'
            }
        
        # Check if file is not empty
        if file.size == 0:
            return {'is_valid': False, 'error': 'File is empty'}
        
        return {'is_valid': True, 'error': None}



class TextExtractor:
    """Handles text extraction from different file formats"""
    
    def __init__(self):
        self.extraction_methods = {
            'pdf': self._extract_from_pdf,
            'docx': self._extract_from_docx,
            'txt': self._extract_from_txt
        }
        self.logger = self._setup_logger()
    
    def extract_text(self, file_info: Dict) -> Dict[str, any]:
        """Main text extraction method"""
        file_format = file_info['format']
        
        try:
            if file_format not in self.extraction_methods:
                raise ValueError(f"Unsupported format: {file_format}")
            
            # Extract text using appropriate method
            extracted_text = self.extraction_methods[file_format](file_info['content'])
            
            # Validate extraction result
            if not extracted_text or len(extracted_text.strip()) < 10:
                raise ValueError("Extracted text is too short or empty")
            
            self.logger.info(f"Successfully extracted text from {file_info['name']}")
            
            return {
                'success': True,
                'text': extracted_text,
                'word_count': len(extracted_text.split()),
                'char_count': len(extracted_text),
                'extraction_method': file_format,
                'file_name': file_info['name']
            }
            
        except Exception as e:
            error_msg = f"Extraction failed for {file_info['name']}: {str(e)}"
            self.logger.error(error_msg)
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'word_count': 0,
                'char_count': 0,
                'file_name': file_info['name']
            }
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF files with fallback methods"""
        text = ""
        
        try:
            # Primary method: PyPDF2
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                except Exception as e:
                    self.logger.warning(f"Failed to extract page {page_num + 1}: {e}")
            
            # If PyPDF2 extraction is poor, try alternative method
            if len(text.strip()) < 50:
                raise ValueError("PyPDF2 extraction yielded insufficient text")
            
        except Exception as e:
            # Fallback method: Try to use textract (if available)
            try:
                import textract
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                    tmp.write(file_content)
                    tmp_path = tmp.name
                
                try:
                    text = textract.process(tmp_path).decode('utf-8')
                finally:
                    os.unlink(tmp_path)
            except ImportError:
                raise Exception("PDF extraction failed. Install textract: pip install textract")
            except Exception as fallback_error:
                raise Exception(f"All PDF extraction methods failed: {str(e)}, {str(fallback_error)}")
        
        return text.strip()
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(BytesIO(file_content))
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            
            if not text.strip():
                raise ValueError("No text content found in DOCX file")
            
            return text
            
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def _extract_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT files with encoding detection"""
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                # Validate that the text makes sense
                if len(text.strip()) > 0:
                    return text
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        # If all encodings fail, use error handling
        try:
            return file_content.decode('utf-8', errors='replace')
        except Exception as e:
            raise Exception(f"TXT extraction failed: {str(e)}")
    
    def _setup_logger(self):
        """Setup logging for extraction process"""
        logger = logging.getLogger('TextExtractor')
        if not logger.handlers:
            logger.setLevel(logging.INFO)
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        return logger


class TextCleaner:
    """Handles text cleaning and preprocessing"""
    
    def __init__(self):
        self.cleaning_steps = [
            self._remove_extra_whitespace,
            self._normalize_line_breaks,
            self._remove_special_characters,
            self._standardize_formatting,
            self._extract_sections
        ]
    
    def clean_text(self, raw_text: str, document_type: str = 'general') -> Dict[str, any]:
        """Main text cleaning pipeline"""
        if not raw_text or not raw_text.strip():
            return {
                'success': False,
                'error': 'Empty or invalid text provided',
                'cleaned_text': '',
                'cleaning_log': []
            }
        
        try:
            cleaned_text = raw_text
            cleaning_log = []
            
            # Apply cleaning steps sequentially
            for step in self.cleaning_steps:
                before_length = len(cleaned_text)
                cleaned_text = step(cleaned_text, document_type)
                after_length = len(cleaned_text)
                
                step_name = step.__name__.replace('_', ' ').title()
                cleaning_log.append({
                    'step': step_name,
                    'chars_before': before_length,
                    'chars_after': after_length,
                    'reduction': before_length - after_length
                })
            
            # Calculate final statistics
            original_length = len(raw_text)
            final_length = len(cleaned_text)
            reduction_percentage = ((original_length - final_length) / original_length * 100) if original_length > 0 else 0
            
            return {
                'success': True,
                'cleaned_text': cleaned_text,
                'original_length': original_length,
                'final_length': final_length,
                'reduction_percentage': reduction_percentage,
                'cleaning_log': cleaning_log
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'cleaned_text': raw_text,
                'cleaning_log': []
            }
    
    def _remove_extra_whitespace(self, text: str, doc_type: str) -> str:
        """Remove extra whitespace and normalize spacing"""
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        
        # Remove leading/trailing whitespace from lines
        lines = [line.strip() for line in text.split('\n')]
        
        # Remove completely empty lines but preserve paragraph breaks
        cleaned_lines = []
        prev_empty = False
        
        for line in lines:
            if line:
                cleaned_lines.append(line)
                prev_empty = False
            elif not prev_empty:  # Keep one empty line for paragraph breaks
                cleaned_lines.append('')
                prev_empty = True
        
        return '\n'.join(cleaned_lines).strip()
    
    def _normalize_line_breaks(self, text: str, doc_type: str) -> str:
        """Normalize different line break formats"""
        # Convert Windows/Mac line breaks to Unix
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive line breaks (more than 2 consecutive)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text
    
    def _remove_special_characters(self, text: str, doc_type: str) -> str:
        """Remove unwanted special characters while preserving structure"""
        # Remove common PDF artifacts and control characters
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
        
        # Normalize bullet points
        text = re.sub(r'[•·▪▫■□◦‣⁃→‰]', '• ', text)
        
        # Normalize quotes
        text = re.sub(r'[""''«»]', '"', text)
        
        # Remove common document artifacts
        text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'^--- Page \d+ ---$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\d+$', '', text, flags=re.MULTILINE)  # Remove standalone numbers
        
        return text
    
    def _standardize_formatting(self, text: str, doc_type: str) -> str:
        """Standardize common formatting patterns"""
        # Standardize email format (but keep readable)
        text = re.sub(
            r'\b([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})\b', 
            r'\1', text
        )
        
        # Standardize phone numbers
        text = re.sub(
            r'\b(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})\b', 
            r'(\1) \2-\3', text
        )
        
        # Normalize common separators
        text = re.sub(r'\s*[|•]\s*', ' | ', text)
        
        return text
    
    def _extract_sections(self, text: str, doc_type: str) -> str:
        """Extract and organize document sections"""
        if doc_type == 'resume':
            return self._process_resume_sections(text)
        elif doc_type == 'job_description':
            return self._process_job_description_sections(text)
        
        return text
    
    def _process_resume_sections(self, text: str) -> str:
        """Process resume-specific sections"""
        # Common resume section headers (case insensitive)
        section_patterns = [
            (r'\b(PROFESSIONAL\s+EXPERIENCE|WORK\s+EXPERIENCE|EXPERIENCE|EMPLOYMENT\s+HISTORY)\b', 'EXPERIENCE'),
            (r'\b(EDUCATION|EDUCATIONAL\s+BACKGROUND|ACADEMIC\s+BACKGROUND)\b', 'EDUCATION'),
            (r'\b(TECHNICAL\s+SKILLS|SKILLS|CORE\s+COMPETENCIES|COMPETENCIES)\b', 'SKILLS'),
            (r'\b(PROJECTS|PROJECT\s+EXPERIENCE|NOTABLE\s+PROJECTS)\b', 'PROJECTS'),
            (r'\b(CERTIFICATIONS|LICENSES|PROFESSIONAL\s+CERTIFICATIONS)\b', 'CERTIFICATIONS'),
            (r'\b(ACHIEVEMENTS|ACCOMPLISHMENTS|AWARDS)\b', 'ACHIEVEMENTS')
        ]
        
        for pattern, section_name in section_patterns:
            text = re.sub(pattern, f'\n=== {section_name} ===', text, flags=re.IGNORECASE)
        
        return text
    
    def _process_job_description_sections(self, text: str) -> str:
        """Process job description specific sections"""
        jd_patterns = [
            (r'\b(REQUIREMENTS|REQUIRED\s+QUALIFICATIONS|MINIMUM\s+REQUIREMENTS)\b', 'REQUIREMENTS'),
            (r'\b(RESPONSIBILITIES|JOB\s+DUTIES|KEY\s+RESPONSIBILITIES|DUTIES)\b', 'RESPONSIBILITIES'),
            (r'\b(PREFERRED|NICE\s+TO\s+HAVE|PREFERRED\s+QUALIFICATIONS)\b', 'PREFERRED'),
            (r'\b(BENEFITS|COMPENSATION|SALARY|PACKAGE)\b', 'BENEFITS'),
            (r'\b(ABOUT|COMPANY|ORGANIZATION|OVERVIEW)\b', 'ABOUT')
        ]
        
        for pattern, section_name in jd_patterns:
            text = re.sub(pattern, f'\n=== {section_name} ===', text, flags=re.IGNORECASE)
        
        return text


class DocumentProcessor:
    """Main document processing pipeline coordinator"""
    
    def __init__(self):
        self.uploader = DocumentUploader()
        self.extractor = TextExtractor()
        self.cleaner = TextCleaner()
        self.processed_documents = []
    
    def run_pipeline(self):
        """Execute the complete document processing pipeline"""
        # Initialize session state
        if 'processed_docs' not in st.session_state:
            st.session_state.processed_docs = []
        
        # Step 1: Upload Interface
        uploaded_files = self.uploader.create_upload_interface()
        
        if uploaded_files:
            # Step 2: Process Documents Button
            col1, col2, col3 = st.columns([1, 1, 1])
            with col2:
                if st.button("🚀 Click Process Documents", type="primary", use_container_width=True):
                    self._process_all_documents(uploaded_files)
        
        # Step 3: Display Results (if any processed documents exist)
        if st.session_state.processed_docs:
            self._display_processing_results()
            self._create_download_interface()
    
    def _process_all_documents(self, uploaded_files: List[Dict]):
        """Process all uploaded documents"""
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        processed_docs = []
        total_files = len(uploaded_files)
        
        for i, file_info in enumerate(uploaded_files):
            # Update progress
            progress = (i + 1) / total_files
            progress_bar.progress(progress)
            status_text.text(f"Processing {file_info['name']} ({i+1}/{total_files})")
            
            # Process single document
            result = self._process_single_document(file_info)
            processed_docs.append(result)
        
        # Store in session state
        st.session_state.processed_docs = processed_docs
        
        # Clear progress indicators
        progress_bar.empty()
        status_text.empty()
        
        # Show completion message
        successful = sum(1 for doc in processed_docs if doc['success'])
        st.success(f"✅ Processing complete! {successful}/{total_files} documents processed successfully.")
    
    def _process_single_document(self, file_info: Dict) -> Dict:
        """Process a single document through the complete pipeline"""
        try:
            # Step 1: Extract text
            extraction_result = self.extractor.extract_text(file_info)
            
            if not extraction_result['success']:
                return {
                    'file_name': file_info['name'],
                    'document_type': file_info['type'],
                    'success': False,
                    'error': extraction_result['error'],
                    'stage': 'extraction',
                    'timestamp': datetime.now()
                }
            
            # Step 2: Clean text
            cleaning_result = self.cleaner.clean_text(
                extraction_result['text'], 
                file_info['type']
            )
            
            if not cleaning_result['success']:
                return {
                    'file_name': file_info['name'],
                    'document_type': file_info['type'],
                    'success': False,
                    'error': cleaning_result['error'],
                    'stage': 'cleaning',
                    'timestamp': datetime.now()
                }
            
            # Success case
            return {
                'file_name': file_info['name'],
                'document_type': file_info['type'],
                'success': True,
                'original_text': extraction_result['text'],
                'cleaned_text': cleaning_result['cleaned_text'],
                'extraction_stats': {
                    'word_count': extraction_result['word_count'],
                    'char_count': extraction_result['char_count'],
                    'extraction_method': extraction_result['extraction_method']
                },
                'cleaning_stats': {
                    'original_length': cleaning_result['original_length'],
                    'final_length': cleaning_result['final_length'],
                    'reduction_percentage': cleaning_result['reduction_percentage']
                },
                'processing_log': cleaning_result['cleaning_log'],
                'timestamp': datetime.now()
            }
            
        except Exception as e:
            return {
                'file_name': file_info['name'],
                'document_type': file_info['type'],
                'success': False,
                'error': str(e),
                'stage': 'general_processing',
                'timestamp': datetime.now()
            }
    
    def _display_processing_results(self):
        """Display comprehensive processing results"""
        st.header("📊 Processing Results")
        
        processed_docs = st.session_state.processed_docs
        
        # Summary metrics
        successful = sum(1 for doc in processed_docs if doc['success'])
        total = len(processed_docs)
        success_rate = (successful / total * 100) if total > 0 else 0
        
        # Display metrics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("📄 Total Documents", total)
        with col2:
            st.metric("✅ Successfully Processed", successful)
        with col3:
            st.metric("❌ Failed", total - successful)
        with col4:
            st.metric("🎯 Success Rate", f"{success_rate:.1f}%")
        
        # Detailed results tabs
        success_docs = [doc for doc in processed_docs if doc['success']]
        failed_docs = [doc for doc in processed_docs if not doc['success']]
        
        if success_docs:
            success_tab, failed_tab = st.tabs([f"✅ Successful ({len(success_docs)})", f"❌ Failed ({len(failed_docs)})"])
        else:
            failed_tab = st.tabs([f"❌ Failed ({len(failed_docs)})"])[0]
        
        # Display successful documents
        if success_docs:
            with success_tab:
                for i, doc in enumerate(success_docs):
                    self._display_successful_document(doc, i)
        
        # Display failed documents
        if failed_docs:
            with failed_tab:
                for doc in failed_docs:
                    self._display_failed_document(doc)
    
    def _display_successful_document(self, doc: Dict, index: int):
        """Display details for a successfully processed document"""
        with st.expander(f"📄 {doc['file_name']} ({doc['document_type']})"):
            # Statistics columns
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("📈 Extraction Statistics")
                st.write(f"**Words:** {doc['extraction_stats']['word_count']:,}")
                st.write(f"**Characters:** {doc['extraction_stats']['char_count']:,}")
                st.write(f"**Method:** {doc['extraction_stats']['extraction_method'].upper()}")
            
            with col2:
                st.subheader("🧹 Cleaning Statistics")
                st.write(f"**Original Length:** {doc['cleaning_stats']['original_length']:,} chars")
                st.write(f"**Final Length:** {doc['cleaning_stats']['final_length']:,} chars")
                st.write(f"**Reduction:** {doc['cleaning_stats']['reduction_percentage']:.1f}%")
            
            # Processing log
            if doc['processing_log']:
                st.subheader("🔍 Processing Steps")
                log_df = pd.DataFrame(doc['processing_log'])
                st.dataframe(log_df, use_container_width=True)
            
            # Text preview
            st.subheader("📖 Text Preview")
            preview_text = doc['cleaned_text'][:1000] + "..." if len(doc['cleaned_text']) > 1000 else doc['cleaned_text']
            st.text_area(
                "Cleaned Text",
                preview_text,
                height=300,
                key=f"preview_{index}",
                disabled=True
            )
    
    def _display_failed_document(self, doc: Dict):
        """Display details for a failed document"""
        with st.expander(f"❌ {doc['file_name']} - Failed at {doc['stage']}"):
            st.error(f"**Error:** {doc['error']}")
            st.write(f"**Document Type:** {doc['document_type']}")
            st.write(f"**Processing Stage:** {doc['stage']}")
            st.write(f"**Timestamp:** {doc['timestamp'].strftime('%Y-%m-%d %H:%M:%S')}")
    
    def _create_download_interface(self):
        """Create download interface for processed documents"""
        successful_docs = [doc for doc in st.session_state.processed_docs if doc['success']]
        
        if not successful_docs:
            return
        
        st.header("💾 Download Processed Data")
        
        # Download options
        col1, col2 = st.columns(2)
        
        with col1:
            # CSV download
            csv_data = self._create_csv_export(successful_docs)
            st.download_button(
                label="📥 Download as CSV",
                data=csv_data,
                file_name=f"processed_documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
                use_container_width=True
            )
        
        with col2:
            # JSON download
            json_data = self._create_json_export(successful_docs)
            st.download_button(
                label="📥 Download as JSON",
                data=json_data,
                file_name=f"processed_documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
                use_container_width=True
            )
    
    def _create_csv_export(self, documents: List[Dict]) -> str:
        """Create CSV export of processed documents"""
        export_data = []
        
        for doc in documents:
            export_data.append({
                'filename': doc['file_name'],
                'document_type': doc['document_type'],
                'word_count': doc['extraction_stats']['word_count'],
                'char_count': doc['extraction_stats']['char_count'],
                'original_length': doc['cleaning_stats']['original_length'],
                'final_length': doc['cleaning_stats']['final_length'],
                'reduction_percentage': doc['cleaning_stats']['reduction_percentage'],
                'cleaned_text': doc['cleaned_text'],
                'processing_timestamp': doc['timestamp'].strftime('%Y-%m-%d %H:%M:%S')
            })
        
        df = pd.DataFrame(export_data)
        return df.to_csv(index=False)
    
    def _create_json_export(self, documents: List[Dict]) -> str:
        """Create JSON export of processed documents"""
        export_data = []
        
        for doc in documents:
            # Create a clean export without internal objects
            clean_doc = {
                'filename': doc['file_name'],
                'document_type': doc['document_type'],
                'cleaned_text': doc['cleaned_text'],
                'extraction_stats': doc['extraction_stats'],
                'cleaning_stats': doc['cleaning_stats'],
                'processing_timestamp': doc['timestamp'].isoformat()
            }
            export_data.append(clean_doc)
        
        return json.dumps(export_data, indent=2)


def main():
    """Main application entry point"""
    try:
        # Initialize and run the document processor
        processor = DocumentProcessor()
        processor.run_pipeline()
        
        # Sidebar information
        with st.sidebar:
            st.header("ℹ️ Information")
            st.markdown("""
            **Supported Formats:**
            - PDF files
            - DOCX files  
            - TXT files
            
            **File Size Limit:** 10MB per file
            
            **Processing Steps:**
            1. File upload & validation
            2. Text extraction
            3. Text cleaning & preprocessing
            4. Results display & download
            """)
            
            st.header("📋 Status")
            if st.session_state.get('processed_docs'):
                total = len(st.session_state.processed_docs)
                successful = sum(1 for doc in st.session_state.processed_docs if doc['success'])
                st.success(f"Processed: {successful}/{total} documents")
            else:
                st.info("No documents processed yet")
    
    except Exception as e:
        st.error(f"Application error: {str(e)}")
        st.error("Please refresh the page and try again.")


if __name__ == "__main__":
    main()